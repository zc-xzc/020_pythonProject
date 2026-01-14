import os
import re
import shutil
import random
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime, timedelta
from docx import Document
import PyPDF2
from collections import defaultdict, Counter
import matplotlib
import sys  # 用于处理命令行参数

matplotlib.use('Agg')  # 非交互式环境下使用
# 修正字体配置，只保留系统常用中文字体
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]  # 支持中文显示
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

# -------------------------- 多维度分类标准配置 --------------------------
LEVELS = {
    '国家层面': ['国家', '国务院', '中央', '全国', '人大', '部委', '宏观导向', '顶层设计', '国家战略',
                 '政治局', '常委会', '最高人民法院', '最高人民检察院'],
    '省层面': ['省', '自治区', '直辖市', '省人民政府', '省级', '省内', '省级规划', '省委', '省级部门'],
    '市层面': ['市', '自治州', '市人民政府', '市级', '市内', '市级规划', '市委', '市级部门'],
    '区/县层面': ['区', '县', '区政府', '县政府', '区级', '县级', '区县规划']
}

DOC_TYPES = {
    '法律': ['法', '人大颁布', '主席令', '法典', '法律'],
    '行政法规': ['条例', '行政法规', '国务院令', '实施条例'],
    '地方性法规': ['地方性法规', '省人大', '市人大', '自治区人大'],
    '规章': ['规定', '办法', '细则', '规章', '部委令', '地方政府令'],
    '规划': ['规划', '战略', '纲要', '中长期目标', '发展纲要', '专项规划'],
    '标准': ['标准', '规范', '技术要求', '排放标准', '排放限值', '技术导则', 'GB/T', 'HJ'],
    '政策文件': ['通知', '意见', '决定', '指示', '批复', '函', '政策'],
    '执法文件': ['执法', '通报', '处罚决定', '督查结果', '检查结果', '监察报告'],
    '行动方案': ['行动方案', '实施方案', '攻坚计划', '专项行动', '工作计划']
}

ENVIRONMENTAL_FACTORS = {
    '大气': ['大气', '空气', 'PM2.5', 'PM10', '雾霾', '扬尘', '挥发性有机物', 'VOCs', '二氧化硫', '氮氧化物'],
    '水': ['水', '水质', '河流', '湖泊', '黑臭水体', '污水处理', '水源地', '地下水', '地表水', '水环境'],
    '土壤': ['土壤', '土地', '重金属', '土壤修复', '耕地污染', '农用地', '建设用地', '土壤质量'],
    '碳排放': ['碳', '碳中和', '碳达峰', '双碳', '低碳', '碳交易', '碳市场', '温室气体', '碳排放权'],
    '污染源': ['污染源', '排污', '污染物', '污染源普查', '工业污染', '农业污染', '生活污染', '面源污染'],
    '生态保护': ['生态', '生态红线', '自然保护区', '生物多样性', '生态修复', '生态系统', '湿地', '森林'],
    '固废与危废': ['固体废物', '垃圾', '危废', '危险废物', '固废', '垃圾分类', '废物处理', '无废城市'],
    '噪声与振动': ['噪声', '振动', '声环境', '噪音污染', '声学'],
    '辐射': ['辐射', '放射性', '电磁辐射', '电离辐射', '核安全']
}

POLICY_TOOLS = {
    '命令控制型': ['强制', '必须', '禁止', '处罚', '关停', '限期整改', '强制执行', '标准', '许可', '审批'],
    '市场激励型': ['补贴', '奖励', '税收优惠', '碳交易', '排污权交易', '绿色金融', '环境税', '生态补偿', '绿色信贷'],
    '自愿参与型': ['自愿', '倡议', '承诺', '认证', '企业自律', '行业公约', '环境标志', '绿色认证'],
    '多元协同型': ['公众监督', 'NGO', '社会组织', '跨区域合作', '政企联动', '公众参与', '信息公开', '社会共治'],
    '技术支持型': ['技术推广', '研发', '科技', '创新', '技术标准', '示范工程', '技术改造']
}

REGIONS = {
    '京津冀': ['北京', '天津', '河北', '京津冀', '雄安', '张家口', '承德'],
    '长三角': ['上海', '江苏', '浙江', '安徽', '长三角', '沪苏浙皖', '南京', '杭州'],
    '珠三角': ['广东', '广州', '深圳', '珠三角', '粤港澳', '大湾区', '珠海', '东莞'],
    '东北地区': ['辽宁', '吉林', '黑龙江', '东北', '沈阳', '长春', '哈尔滨'],
    '中部地区': ['山西', '河南', '湖北', '湖南', '江西', '安徽', '中部'],
    '西部地区': ['重庆', '四川', '贵州', '云南', '西藏', '陕西', '甘肃', '青海', '宁夏', '新疆', '西部'],
    '成渝地区': ['重庆', '成都', '成渝', '四川东部', '重庆西部'],
    '长江经济带': ['长江', '长江经济带', '沿江省市', '长江流域'],
    '黄河流域': ['黄河', '黄河流域', '黄河生态经济带'],
    '全国性': ['全国', '无特定区域', '各地区', '所有地区']
}

CORE_KEYWORDS = [
    '污染防治', '排放标准', '生态保护', '碳达峰', '碳中和', '绿色转型',
    '环境治理', '执法检查', '惩处机制', '激励补贴', '双碳', '低碳转型',
    '生态文明', '可持续发展', '环境监测', '环境评估', '生态补偿',
    '绿色金融', '清洁生产', '循环经济', '无废城市', '环境修复',
    '环境规划', '环境标准', '排污许可', '环境影响评价'
]

START_YEAR = 2015
END_YEAR = 2025
START_DATE = datetime(2015, 1, 1)
END_DATE = datetime(2025, 12, 31)


# -------------------------- 文件读取与内容解析 --------------------------
def read_file_content(file_path):
    """读取txt/docx/pdf内容，支持多编码和错误处理"""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.txt':
            return read_txt(file_path)
        elif ext == '.docx':
            return read_docx(file_path)
        elif ext == '.pdf':
            return read_pdf(file_path)
        else:
            print(f"不支持的文件类型: {ext}（{file_path}）")
            return ""
    except Exception as e:
        print(f"读取失败 {file_path}：{str(e)}")
        return ""


def read_txt(file_path):
    """尝试多种编码读取txt，提高兼容性"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'ansi', 'utf-16']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    print(f"无法解析编码：{file_path}")
    return ""


def read_docx(file_path):
    """读取docx文件文本内容，包括页眉页脚"""
    doc = Document(file_path)
    full_text = []

    # 读取正文
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 读取页眉
    for section in doc.sections:
        for header in section.headers:
            for para in header.paragraphs:
                full_text.append(para.text)

    # 读取页脚
    for section in doc.sections:
        for footer in section.footers:
            for para in footer.paragraphs:
                full_text.append(para.text)

    return ' '.join(full_text)


def read_pdf(file_path):
    """读取pdf文件文本内容（支持多页和加密文件）"""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            # 检查是否加密
            if reader.is_encrypted:
                try:
                    # 尝试空密码解密
                    reader.decrypt('')
                except:
                    print(f"加密PDF文件无法读取：{file_path}")
                    return ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
        return text
    except Exception as e:
        print(f"PDF读取错误 {file_path}：{str(e)}")
        return ""


# -------------------------- 核心信息提取与分类 --------------------------
def is_relevant(text):
    """判断文件是否与核心主题相关（至少匹配1个核心关键词）"""
    text_lower = text.lower()
    return any(kw.lower() in text_lower for kw in CORE_KEYWORDS)


def extract_date(text, file_name):
    """提取文件日期（支持多种格式，优先文本后文件名）"""
    date_patterns = [
        r'\b20\d{2}[/-]\d{1,2}[/-]\d{1,2}\b',  # YYYY-MM-DD / YYYY/MM/DD
        r'\b20\d{2}年\d{1,2}月\d{1,2}日\b',  # YYYY年MM月DD日
        r'\b20\d{2}年\d{1,2}月\b',  # YYYY年MM月
        r'\b20\d{2}年\b'  # YYYY年
    ]

    # 从文本提取
    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            return parse_date(match.group())

    # 从文件名提取
    for pattern in date_patterns:
        match = re.search(pattern, file_name)
        if match:
            return parse_date(match.group())

    return None


def parse_date(date_str):
    """解析日期字符串为datetime对象，支持部分日期"""
    try:
        if '-' in date_str or '/' in date_str:
            date_str = date_str.replace('/', '-')
            # 尝试完整日期
            try:
                return datetime.strptime(date_str, '%Y-%m-%d')
            except:
                # 尝试年月
                try:
                    return datetime.strptime(date_str, '%Y-%m')
                except:
                    # 尝试年份
                    return datetime.strptime(date_str, '%Y')
        elif '年' in date_str:
            # 尝试完整日期
            try:
                return datetime.strptime(date_str, '%Y年%m月%d日')
            except:
                # 尝试年月
                try:
                    return datetime.strptime(date_str, '%Y年%m月')
                except:
                    # 尝试年份
                    return datetime.strptime(date_str, '%Y年')
    except:
        return None
    return None


def extract_year(date_obj):
    """从日期对象提取年份（兼容None值）"""
    if isinstance(date_obj, datetime):
        return date_obj.year if (START_YEAR <= date_obj.year <= END_YEAR) else None
    return None


def classify_by_dimension(text, dimension_dict, priority=False):
    """按维度分类（支持优先级模式和多标签）"""
    matched = []
    text_lower = text.lower()
    for category, keywords in dimension_dict.items():
        if any(kw.lower() in text_lower for kw in keywords):
            matched.append(category)
            if priority:  # 优先级模式：匹配到第一个即返回
                return matched
    return matched if matched else ['未明确']


def get_region(text, file_name):
    """区域分类（优先文件内容，补充文件名，支持多区域）"""
    combined_text = (text + " " + file_name).lower()
    regions = []
    for region, keywords in REGIONS.items():
        if any(kw.lower() in combined_text for kw in keywords):
            regions.append(region)

    # 去重并处理全国性与其他区域的关系
    regions = list(set(regions))
    if '全国性' in regions and len(regions) > 1:
        regions.remove('全国性')

    return regions if regions else ['全国性']


# -------------------------- 单文件分析与批量处理 --------------------------
def analyze_file(file_path):
    """全维度分析单个文件，返回结构化信息"""
    file_name = os.path.basename(file_path)
    text = read_file_content(file_path)
    if not text:
        return None

    # 核心信息提取
    date = extract_date(text, file_name)
    year = extract_year(date)
    in_time_range = "是" if (date and START_DATE <= date <= END_DATE) else "否"
    relevant = is_relevant(text)

    # 多维度分类 - 确保有默认值
    level = classify_by_dimension(text, LEVELS, priority=True) or ['未明确']
    doc_type = classify_by_dimension(text, DOC_TYPES) or ['未明确']
    env_factor = classify_by_dimension(text, ENVIRONMENTAL_FACTORS) or ['未明确']
    policy_tool = classify_by_dimension(text, POLICY_TOOLS) or ['未明确']
    region = get_region(text, file_name) or ['全国性']
    keywords = extract_keywords(text)

    return {
        '文件名': file_name,
        '文件路径': file_path,
        '提取日期': date.strftime('%Y-%m-%d') if date else '未提取到',
        '是否在时间范围': in_time_range,
        '年份': year if year else '未明确',
        '是否相关': '是' if relevant else '否',
        '政府层级': ';'.join(level),
        '文件类型': ';'.join(doc_type),
        '环境要素': ';'.join(env_factor),
        '政策工具类型': ';'.join(policy_tool),
        '所属区域': ';'.join(region),
        '核心关键词': ';'.join(keywords)
    }


def extract_keywords(text):
    """从文本中提取核心关键词"""
    text_lower = text.lower()
    extracted = []
    for kw in CORE_KEYWORDS:
        if kw.lower() in text_lower and kw not in extracted:
            extracted.append(kw)
    return extracted if extracted else ['未匹配关键词']


def batch_analyze(input_dir):
    """批量处理目录下所有文件，返回分析结果列表"""
    results = []
    for root, _, files in os.walk(input_dir):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in ['.txt', '.docx', '.pdf']:
                file_path = os.path.join(root, file)
                info = analyze_file(file_path)
                if info:
                    results.append(info)
    return results


# -------------------------- 交叉分析与报告生成 --------------------------
def generate_cross_analysis(df, output_dir):
    """生成多维度交叉分析报告和图表"""
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 确保基础列存在
    required_columns = ['政府层级', '文件类型', '环境要素', '政策工具类型', '所属区域']
    for col in required_columns:
        if col not in df.columns:
            df[col] = '未明确'  # 为缺失列添加默认值

    # 处理多标签数据，创建主要标签列
    df['主要政府层级'] = df['政府层级'].apply(lambda x: x.split(';')[0] if x else '未明确')
    df['主要文件类型'] = df['文件类型'].apply(lambda x: x.split(';')[0] if x else '未明确')
    df['主要环境要素'] = df['环境要素'].apply(lambda x: x.split(';')[0] if x else '未明确')
    df['主要政策工具'] = df['政策工具类型'].apply(lambda x: x.split(';')[0] if x else '未明确')
    df['主要区域'] = df['所属区域'].apply(lambda x: x.split(';')[0] if x else '未明确')

    # 保存处理后的DataFrame
    df.to_excel(os.path.join(output_dir, '完整分析数据.xlsx'), index=False, engine='openpyxl')

    # 生成分析报告
    report = []
    total_files = len(df)
    relevant_files = len(df[df['是否相关'] == '是'])

    report.append("===== 环境政策文件多维度分析报告 =====")
    report.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append(f"总文件数：{total_files} 份")
    report.append(f"相关文件数：{relevant_files} 份（占比 {relevant_files / total_files:.2%}）")
    report.append(f"时间范围：{START_YEAR}年 - {END_YEAR}年\n")

    # 1. 环境要素 × 政策工具交叉分析
    report.append("----- 1. 环境要素 × 政策工具分布 -----")
    env_tool_cross = pd.crosstab(df['主要环境要素'], df['主要政策工具'], normalize='index')
    report.append(env_tool_cross.to_string())

    # 识别异常组合（占比<5%）
    异常组合 = []
    for env, row in env_tool_cross.iterrows():
        for tool, ratio in row.items():
            if ratio < 0.05 and env != '未明确' and tool != '未明确':
                异常组合.append(f"{env} + {tool}: {ratio:.2%}")

    if 异常组合:
        report.append("\n异常组合（占比<5%）：")
        report.append("\n".join(异常组合))
        report.append("提示：这些组合占比过低，可能存在分类逻辑问题或文献数量不足\n")
    else:
        report.append("\n未发现异常组合\n")

    # 绘制环境要素×政策工具热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(env_tool_cross, annot=True, cmap="YlGnBu", fmt='.2%')
    plt.title('环境要素×政策工具分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '环境要素_政策工具热图.png'))
    plt.close()

    # 2. 政府层级 × 区域交叉分析
    report.append("----- 2. 政府层级 × 区域分布 -----")
    level_region_cross = pd.crosstab(df['主要政府层级'], df['主要区域'], normalize='index')
    report.append(level_region_cross.to_string())

    # 识别异常组合
    异常组合 = []
    for level, row in level_region_cross.iterrows():
        for region, ratio in row.items():
            if ratio < 0.05 and level != '未明确' and region != '全国性' and region != '未明确':
                异常组合.append(f"{level} + {region}: {ratio:.2%}")

    if 异常组合:
        report.append("\n异常组合（占比<5%）：")
        report.append("\n".join(异常组合))
        report.append("提示：这些组合占比过低，可能存在区域政策覆盖不均衡问题\n")
    else:
        report.append("\n未发现异常组合\n")

    # 绘制政府层级×区域热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(level_region_cross, annot=True, cmap="YlOrRd", fmt='.2%')
    plt.title('政府层级×区域分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '政府层级_区域热图.png'))
    plt.close()

    # 3. 文件类型 × 政策工具交叉分析
    report.append("----- 3. 文件类型 × 政策工具分布 -----")
    type_tool_cross = pd.crosstab(df['主要文件类型'], df['主要政策工具'], normalize='index')
    report.append(type_tool_cross.to_string())

    # 绘制文件类型×政策工具热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(type_tool_cross, annot=True, cmap="PuBuGn", fmt='.2%')
    plt.title('文件类型×政策工具分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '文件类型_政策工具热图.png'))
    plt.close()

    # 4. 时间分布分析
    report.append("\n----- 4. 时间分布分析 -----")
    year_distribution = df['年份'].value_counts().sort_index()
    report.append("各年份文件数量：")
    report.append(year_distribution.to_string())

    # 绘制时间分布折线图
    plt.figure(figsize=(12, 6))
    year_distribution.plot(kind='line', marker='o')
    plt.title('文件数量年度分布')
    plt.xlabel('年份')
    plt.ylabel('文件数量')
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '文件数量年度分布.png'))
    plt.close()

    # 5. 各维度分布统计
    report.append("\n----- 5. 各维度分布统计 -----")

    # 政府层级分布 - 修正百分比格式
    level_dist = df['主要政府层级'].value_counts(normalize=True)
    report.append("政府层级分布：")
    report.append(level_dist.apply(lambda x: f"{x:.2%}").to_string())

    # 文件类型分布 - 修正百分比格式
    type_dist = df['主要文件类型'].value_counts(normalize=True)
    report.append("\n文件类型分布：")
    report.append(type_dist.apply(lambda x: f"{x:.2%}").to_string())

    # 环境要素分布 - 修正百分比格式
    env_dist = df['主要环境要素'].value_counts(normalize=True)
    report.append("\n环境要素分布：")
    report.append(env_dist.apply(lambda x: f"{x:.2%}").to_string())

    # 政策工具分布 - 修正百分比格式
    tool_dist = df['主要政策工具'].value_counts(normalize=True)
    report.append("\n政策工具分布：")
    report.append(tool_dist.apply(lambda x: f"{x:.2%}").to_string())

    # 区域分布 - 修正百分比格式
    region_dist = df['主要区域'].value_counts(normalize=True)
    report.append("\n区域分布：")
    report.append(region_dist.apply(lambda x: f"{x:.2%}").to_string())

    # 保存报告文本
    with open(os.path.join(output_dir, '分析报告.txt'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))

    return df


# -------------------------- 准确性验证机制 --------------------------
def generate_sampling_verification(results, output_dir, sample_ratio=0.15):
    """生成随机抽样验证表，用于人工验证"""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 确保样本量合理
    sample_size = max(5, min(50, int(len(results) * sample_ratio)))
    if len(results) < 5:
        sample_size = len(results)

    # 随机抽样
    sampled = random.sample(results, sample_size)

    # 准备验证表格
    verification_data = []
    for i, item in enumerate(sampled, 1):
        verification_data.append({
            '序号': i,
            '文件名': item['文件名'],
            '文件路径': item['文件路径'],
            '机器分类-政府层级': item['政府层级'],
            '人工分类-政府层级': '',
            '机器分类-文件类型': item['文件类型'],
            '人工分类-文件类型': '',
            '机器分类-环境要素': item['环境要素'],
            '人工分类-环境要素': '',
            '机器分类-政策工具': item['政策工具类型'],
            '人工分类-政策工具': '',
            '机器分类-区域': item['所属区域'],
            '人工分类-区域': '',
            '是否一致': '',
            '备注': ''
        })

    # 保存为Excel
    df = pd.DataFrame(verification_data)
    verification_path = os.path.join(output_dir, '人工抽样验证表.xlsx')
    df.to_excel(verification_path, index=False, engine='openpyxl')
    print(f"已生成 {sample_size} 个样本的人工验证表：{verification_path}")

    return verification_path


def calculate_accuracy(verification_file):
    """计算人工验证的准确率"""
    try:
        df = pd.read_excel(verification_file)

        # 检查必要的列是否存在
        required_columns = [
            '机器分类-政府层级', '人工分类-政府层级',
            '机器分类-文件类型', '人工分类-文件类型',
            '机器分类-环境要素', '人工分类-环境要素',
            '机器分类-政策工具', '人工分类-政策工具',
            '机器分类-区域', '人工分类-区域'
        ]

        for col in required_columns:
            if col not in df.columns:
                raise ValueError(f"验证表缺少必要的列：{col}")

        # 计算每个维度的准确率
        total = len(df)
        accurate = {
            '政府层级': 0,
            '文件类型': 0,
            '环境要素': 0,
            '政策工具': 0,
            '区域': 0
        }

        for _, row in df.iterrows():
            # 政府层级准确率（只要有一个匹配就算准确）
            machine_levels = set(row['机器分类-政府层级'].split(';'))
            human_levels = set(str(row['人工分类-政府层级']).split(';'))
            if machine_levels & human_levels:
                accurate['政府层级'] += 1

            # 文件类型准确率
            machine_types = set(row['机器分类-文件类型'].split(';'))
            human_types = set(str(row['人工分类-文件类型']).split(';'))
            if machine_types & human_types:
                accurate['文件类型'] += 1

            # 环境要素准确率
            machine_env = set(row['机器分类-环境要素'].split(';'))
            human_env = set(str(row['人工分类-环境要素']).split(';'))
            if machine_env & human_env:
                accurate['环境要素'] += 1

            # 政策工具准确率
            machine_tool = set(row['机器分类-政策工具'].split(';'))
            human_tool = set(str(row['人工分类-政策工具']).split(';'))
            if machine_tool & human_tool:
                accurate['政策工具'] += 1

            # 区域准确率
            machine_region = set(row['机器分类-区域'].split(';'))
            human_region = set(str(row['人工分类-区域']).split(';'))
            if machine_region & human_region:
                accurate['区域'] += 1

        # 计算总体准确率
        overall_accuracy = sum(accurate.values()) / (total * 5)

        # 输出结果
        print("\n===== 分类准确率报告 =====")
        print(f"验证样本量：{total} 份文件")
        for dimension, count in accurate.items():
            print(f"{dimension}准确率：{count / total:.2%}（{count}/{total}）")
        print(f"总体准确率：{overall_accuracy:.2%}")

        # 保存结果到文件
        with open(os.path.join(os.path.dirname(verification_file), '准确率报告.txt'), 'w', encoding='utf-8') as f:
            f.write("===== 分类准确率报告 =====\n")
            f.write(f"计算时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")
            f.write(f"验证样本量：{total} 份文件\n")
            for dimension, count in accurate.items():
                f.write(f"{dimension}准确率：{count / total:.2%}（{count}/{total}）\n")
            f.write(f"总体准确率：{overall_accuracy:.2%}\n")

        return accurate, overall_accuracy

    except Exception as e:
        print(f"计算准确率失败：{str(e)}")
        return None, None


# -------------------------- 文献缺口分析 --------------------------
def analyze_literature_gaps(df, output_dir):
    """分析文献检索的不足，识别可能的缺口"""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 筛选相关文件
    relevant_df = df[df['是否相关'] == '是'].copy()
    if len(relevant_df) == 0:
        print("没有相关文件可分析文献缺口")
        return

    report = []
    report.append("===== 文献检索不足分析报告 =====")
    report.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append(f"分析基础：{len(relevant_df)} 份相关文件\n")

    # 1. 环境要素覆盖不足分析
    report.append("----- 1. 环境要素覆盖分析 -----")
    env_counts = relevant_df['主要环境要素'].value_counts()
    env_total = sum(env_counts.values)
    env_ratio = env_counts / env_total

    # 识别占比过低的环境要素（<5%）
    low_env = [env for env, ratio in env_ratio.items()
               if ratio < 0.05 and env != '未明确']

    if low_env:
        report.append("环境要素覆盖不均衡，以下要素相关文献占比过低（<5%）：")
        for env in low_env:
            report.append(f"- {env}：{env_counts[env]} 份文件（{env_ratio[env]:.2%}）")
    else:
        report.append("环境要素覆盖较为均衡，未发现明显缺口")

    # 2. 政策工具单一性分析
    report.append("\n----- 2. 政策工具多样性分析 -----")
    tool_counts = relevant_df['主要政策工具'].value_counts()
    tool_total = sum(tool_counts.values)
    tool_ratio = tool_counts / tool_total

    # 识别占比过低的政策工具（<5%）
    low_tool = [tool for tool, ratio in tool_ratio.items()
                if ratio < 0.05 and tool not in ['未明确', '未分类-政策工具']]

    if low_tool:
        report.append("政策工具单一，以下政策工具相关文献占比过低（<5%）：")
        for tool in low_tool:
            report.append(f"- {tool}：{tool_counts[tool]} 份文件（{tool_ratio[tool]:.2%}）")
            # 针对不同政策工具提出补充建议
            suggestions = {
                '市场激励型': '建议补充检索绿色金融、环境税、生态补偿机制等相关政策',
                '自愿参与型': '建议补充检索行业自律规范、环境认证、企业环境承诺等相关政策',
                '多元协同型': '建议补充检索公众参与、社会组织合作、跨区域协同治理等相关政策',
                '技术支持型': '建议补充检索环境技术推广、环保科技研发、示范工程等相关政策'
            }
            report.append(f"  补充建议：{suggestions.get(tool, '建议补充该类型政策工具相关文件')}")
    else:
        report.append("政策工具分布较为均衡，未发现明显单一性问题")

    # 区域覆盖不全分析
    report.append("\n----- 3. 区域覆盖分析 -----")
    region_counts = relevant_df['主要区域'].value_counts()
    region_total = sum(region_counts.values)
    region_ratio = region_counts / region_total

    # 识别占比过低的区域（<3%）
    low_region = [region for region, ratio in region_ratio.items()
                  if ratio < 0.03 and region not in ['未明确', '全国性']]

    if low_region:
        report.append("区域覆盖不足，以下区域相关文献占比过低（<3%）：")
        for region in low_region:
            report.append(f"- {region}：{region_counts[region]} 份文件（{region_ratio[region]:.2%}）")
    else:
        report.append("区域覆盖较为全面，未发现明显缺口")

    # 时间分布不均分析
    report.append("\n----- 4. 时间分布分析 -----")
    year_counts = relevant_df['年份'].value_counts()
    all_years = set(range(START_YEAR, END_YEAR + 1))
    covered_years = set([int(y) for y in year_counts.index if str(y).isdigit()])
    missing_years = all_years - covered_years

    # 近年文献占比（过去3年）
    recent_years = set(range(datetime.now().year - 3, datetime.now().year + 1))
    recent_covered = covered_years & recent_years
    recent_count = sum([year_counts[str(y)] for y in recent_covered if str(y) in year_counts.index])
    recent_ratio = recent_count / len(relevant_df)

    if missing_years:
        report.append(f"时间覆盖不全，缺失以下年份的文献：{', '.join(map(str, sorted(missing_years)))}")

    if recent_ratio < 0.3:
        report.append(f"近年文献不足：近3年文献仅占{recent_ratio:.2%}，可能存在时效性缺口")
    else:
        report.append(f"近年文献充足：近3年文献占比{recent_ratio:.2%}")

    # 文件类型单一性分析
    report.append("\n----- 5. 文件类型多样性分析 -----")
    type_counts = relevant_df['主要文件类型'].value_counts()
    type_total = sum(type_counts.values)
    type_ratio = type_counts / type_total

    # 识别占比过低的文件类型（<5%）
    low_type = [doc_type for doc_type, ratio in type_ratio.items()
                if ratio < 0.05 and doc_type != '未明确']

    if low_type:
        report.append("文件类型单一，以下类型文献占比过低（<5%）：")
        for doc_type in low_type:
            report.append(f"- {doc_type}：{type_counts[doc_type]} 份文件（{type_ratio[doc_type]:.2%}）")
            suggestions = {
                '法律': '建议补充检索环境保护相关法律、人大颁布的环境法规',
                '行政法规': '建议补充检索国务院发布的环境行政法规、实施条例',
                '地方性法规': '建议补充检索地方人大发布的环境法规',
                '标准': '建议补充检索环境质量标准、排放标准、技术规范'
            }
            report.append(f"  补充建议：{suggestions.get(doc_type, '建议补充该类型文件')}")
    else:
        report.append("文件类型分布较为均衡，未发现明显缺失")

    # 6. 主题相关性不足分析
    report.append("\n----- 6. 主题相关性不足分析 -----")

    # 未匹配关键词的文件占比
    irrelevant_count = len(df[df['是否相关'] == '否'])
    irrelevant_ratio = irrelevant_count / len(df)

    if irrelevant_ratio > 0.2:  # 不相关文件占比超过20%
        report.append(f"不相关文件占比过高：{irrelevant_ratio:.2%}的文件被标记为不相关")
        report.append("  补充建议：检查检索策略，可能需要调整关键词或检索范围")

    # 未匹配关键词的情况
    keyword_counts = df['核心关键词'].apply(lambda x: '未匹配关键词' in x).value_counts()
    if True in keyword_counts:
        unmatch_ratio = keyword_counts[True] / len(df)
        if unmatch_ratio > 0.1:  # 未匹配关键词的文件占比超过10%
            report.append(f"\n关键词匹配率低：{unmatch_ratio:.2%}的文件未匹配到核心关键词")
            report.append("  补充建议：扩展核心关键词库，可考虑添加'无废城市'、'清洁生产'、'环境监测'等术语")

    # 7. 综合建议
    report.append("\n----- 7. 综合补充建议 -----")
    report.append("基于以上分析，建议从以下方面补充检索：")

    # 汇总所有建议
    suggestions = []
    if low_env:
        suggestions.append(f"1. 补充{'、'.join(low_env)}等环境要素相关政策文件")
    if low_tool:
        suggestions.append(f"2. 补充{'、'.join(low_tool)}等类型政策工具文件")
    if low_region:
        suggestions.append(f"3. 补充{'、'.join(low_region)}等区域的地方性政策")
    if missing_years or recent_ratio < 0.3:
        suggestions.append(f"4. 补充{START_YEAR}-{END_YEAR}年间缺失或不足年份的政策文件")
    if low_type:
        suggestions.append(f"5. 补充{'、'.join(low_type)}等类型的政策文件")

    if not suggestions:
        report.append("未发现明显的文献缺口，当前文献覆盖较为全面")
    else:
        report.extend(suggestions)

    # 保存报告
    report_path = os.path.join(output_dir, '文献检索不足分析报告.txt')
    with open(report_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))

    print(f"文献检索不足分析完成，报告已保存至：{report_path}")


# -------------------------- 结果排序与缺失资料识别 --------------------------
def sort_and_identify_missing(results, output_dir):
    """按时间和层级排序，并识别缺失资料"""
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    df = pd.DataFrame(results)

    # 确保必要的列存在
    if '政府层级' not in df.columns:
        df['政府层级'] = '未明确'
    if '主要环境要素' not in df.columns:
        df['主要环境要素'] = df['环境要素'].apply(lambda x: x.split(';')[0] if pd.notna(x) else '未明确')
    if '年份' not in df.columns:
        df['年份'] = '未明确'

    # 1. 按时间排序
    time_sorted = df.sort_values(by=['年份', '提取日期'], ascending=False)
    time_sorted_path = os.path.join(output_dir, '按时间排序的文件列表.xlsx')
    time_sorted.to_excel(time_sorted_path, index=False, engine='openpyxl')

    # 2. 按政府层级排序
    level_order = ['国家层面', '省层面', '市层面', '区/县层面', '未明确']
    df['层级排序'] = df['政府层级'].apply(lambda x:
                                          level_order.index(x.split(';')[0]) if (
                                                      pd.notna(x) and x.split(';')[0] in level_order) else len(
                                              level_order))
    level_sorted = df.sort_values(by=['层级排序', '年份'], ascending=[True, False])
    level_sorted_path = os.path.join(output_dir, '按政府层级排序的文件列表.xlsx')
    level_sorted.to_excel(level_sorted_path, index=False, engine='openpyxl')

    # 3. 识别各维度缺失的资料
    missing_report = []
    missing_report.append("===== 缺失资料识别报告 =====")
    missing_report.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    # 按环境要素和年份识别缺失
    missing_report.append("----- 各环境要素按年份的文件分布 -----")
    env_year_pivot = pd.pivot_table(
        df[df['是否相关'] == '是'],
        index='主要环境要素',
        columns='年份',
        values='文件名',
        aggfunc='count',
        fill_value=0
    )
    missing_report.append(env_year_pivot.to_string())

    # 识别各环境要素缺失的年份
    missing_by_env = {}
    for env in env_year_pivot.index:
        if env == '未明确':
            continue
        missing_years = [str(year) for year in env_year_pivot.columns
                         if env_year_pivot.loc[env, year] == 0 and str(year).isdigit()]
        if missing_years:
            missing_by_env[env] = missing_years

    if missing_by_env:
        missing_report.append("\n----- 各环境要素缺失的年份 -----")
        for env, years in missing_by_env.items():
            missing_report.append(f"{env}：{', '.join(years)}")

    # 按政府层级和区域识别缺失
    missing_report.append("\n----- 政府层级与区域的文件分布 -----")
    level_region_pivot = pd.pivot_table(
        df[df['是否相关'] == '是'],
        index='主要政府层级',
        columns='主要区域',
        values='文件名',
        aggfunc='count',
        fill_value=0
    )
    missing_report.append(level_region_pivot.to_string())

    # 识别政府层级与区域的缺失组合
    missing_level_region = []
    for level in level_region_pivot.index:
        if level == '未明确':
            continue
        for region in level_region_pivot.columns:
            if region in ['未明确', '全国性']:
                continue
            if level_region_pivot.loc[level, region] == 0:
                missing_level_region.append(f"{level} + {region}")

    if missing_level_region:
        missing_report.append("\n----- 缺失的政府层级与区域组合 -----")
        missing_report.append(", ".join(missing_level_region))

    # 保存缺失识别报告
    missing_report_path = os.path.join(output_dir, '缺失资料识别报告.txt')
    with open(missing_report_path, 'w', encoding='utf-8') as f:
        f.write('\n'.join(missing_report))

    return time_sorted_path, level_sorted_path


# -------------------------- 主函数 --------------------------
def main(input_dir, output_dir):
    """主函数：协调各模块完成分析流程"""
    print("===== 环境政策文件分析系统 =====")
    print(f"开始处理目录：{input_dir}")

    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 1. 批量分析文件
    print("\n1. 开始批量分析文件...")
    results = batch_analyze(input_dir)
    if not results:
        print("未找到可分析的文件，程序退出")
        return
    print(f"文件分析完成，共处理 {len(results)} 个文件")

    # 2. 生成多维度分析表
    print("\n2. 生成多维度分析表...")
    df = pd.DataFrame(results)
    table_path = os.path.join(output_dir, '政策文件多维度分析表.xlsx')
    df.to_excel(table_path, index=False, engine='openpyxl')
    print(f"多维度分析表已保存至：{table_path}")

    # 3. 生成交叉分析报告和图表
    print("\n3. 生成交叉分析报告和图表...")
    analysis_dir = os.path.join(output_dir, '交叉分析结果')
    analyzed_df = generate_cross_analysis(df, analysis_dir)

    # 4. 生成人工抽样验证表
    print("\n4. 生成人工抽样验证表...")
    verification_dir = os.path.join(output_dir, '人工验证')
    verification_file = generate_sampling_verification(results, verification_dir)

    # 5. 分析文献检索不足
    print("\n5. 分析文献检索不足...")
    gap_dir = os.path.join(output_dir, '文献缺口分析')
    analyze_literature_gaps(analyzed_df, gap_dir)

    # 6. 按时间和层级排序，并识别缺失资料
    print("\n6. 排序文件并识别缺失资料...")
    sort_dir = os.path.join(output_dir, '排序与缺失识别')
    time_sorted_path, level_sorted_path = sort_and_identify_missing(results, sort_dir)

    print("\n===== 分析流程完成 =====")
    print(f"所有结果已保存至：{output_dir}")
    print(f"请查看 {verification_file} 进行人工验证，完成后可运行calculate_accuracy函数计算准确率")


# 执行主函数
if __name__ == "__main__":
    # 示例路径，实际使用时可修改
    input_directory = "D:\Desktop/1020分类：效力位阶 - 重排列"  # 存放待分析文件的目录
    output_directory = "D:\Desktop\新建文件夹"  # 结果输出目录
    main(input_directory, output_directory)

    # 当人工完成验证后，可取消下面的注释计算准确率
    # verification_file = os.path.join(output_directory, '人工验证', '人工抽样验证表.xlsx')
    # calculate_accuracy(verification_file)