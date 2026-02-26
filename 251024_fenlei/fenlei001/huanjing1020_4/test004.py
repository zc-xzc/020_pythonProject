import os
import re
import shutil
import random
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime, timedelta
from docx import Document
import PyPDF2
from collections import defaultdict, Counter
import matplotlib
import sys
from sklearn.metrics import accuracy_score, classification_report

# 设置中文字体和非交互式环境
matplotlib.use('Agg')
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]
plt.rcParams["axes.unicode_minus"] = False  # 解决负号显示问题

# -------------------------- 多维度分类标准配置 --------------------------
LEVELS = {
    '国家层面': ['国家', '国务院', '中央', '全国', '人大', '部委', '宏观导向', '顶层设计', '国家战略',
                 '政治局', '常委会', '最高人民法院', '最高人民检察院'],
    '省层面': ['省', '自治区', '直辖市', '省人民政府', '省级', '省内', '省级规划', '省委', '省级部门'],
    '市层面': ['市', '自治州', '市人民政府', '市级', '市内', '市级规划', '市委', '市级部门'],
    '区/县层面': ['区', '县', '区政府', '县政府', '区级', '县级', '区县规划']
}

DOC_TYPES = {
    '法律': ['法', '人大颁布', '主席令', '法典', '法律'],
    '行政法规': ['条例', '行政法规', '国务院令', '实施条例'],
    '地方性法规': ['地方性法规', '省人大', '市人大', '自治区人大'],
    '规章': ['规定', '办法', '细则', '规章', '部委令', '地方政府令'],
    '规划': ['规划', '战略', '纲要', '中长期目标', '发展纲要', '专项规划'],
    '标准': ['标准', '规范', '技术要求', '排放标准', '排放限值', '技术导则', 'GB/T', 'HJ'],
    '政策文件': ['通知', '意见', '决定', '指示', '批复', '函', '政策'],
    '执法文件': ['执法', '通报', '处罚决定', '督查结果', '检查结果', '监察报告'],
    '行动方案': ['行动方案', '实施方案', '攻坚计划', '专项行动', '工作计划']
}

ENVIRONMENTAL_FACTORS = {
    '大气': ['大气', '空气', 'PM2.5', 'PM10', '雾霾', '扬尘', '挥发性有机物', 'VOCs', '二氧化硫', '氮氧化物'],
    '水': ['水', '水质', '河流', '湖泊', '黑臭水体', '污水处理', '水源地', '地下水', '地表水', '水环境'],
    '土壤': ['土壤', '土地', '重金属', '土壤修复', '耕地污染', '农用地', '建设用地', '土壤质量'],
    '碳排放': ['碳', '碳中和', '碳达峰', '双碳', '低碳', '碳交易', '碳市场', '温室气体', '碳排放权'],
    '污染源': ['污染源', '排污', '污染物', '污染源普查', '工业污染', '农业污染', '生活污染', '面源污染'],
    '生态保护': ['生态', '生态红线', '自然保护区', '生物多样性', '生态修复', '生态系统', '湿地', '森林'],
    '固废与危废': ['固体废物', '垃圾', '危废', '危险废物', '固废', '垃圾分类', '废物处理', '无废城市'],
    '噪声与振动': ['噪声', '振动', '声环境', '噪音污染', '声学'],
    '辐射': ['辐射', '放射性', '电磁辐射', '电离辐射', '核安全']
}

POLICY_TOOLS = {
    '命令控制型': ['强制', '必须', '禁止', '处罚', '关停', '限期整改', '强制执行', '标准', '许可', '审批'],
    '市场激励型': ['补贴', '奖励', '税收优惠', '碳交易', '排污权交易', '绿色金融', '环境税', '生态补偿', '绿色信贷'],
    '自愿参与型': ['自愿', '倡议', '承诺', '认证', '企业自律', '行业公约', '环境标志', '绿色认证'],
    '多元协同型': ['公众监督', 'NGO', '社会组织', '跨区域合作', '政企联动', '公众参与', '信息公开', '社会共治'],
    '技术支持型': ['技术推广', '研发', '科技', '创新', '技术标准', '示范工程', '技术改造']
}

REGIONS = {
    '京津冀': ['北京', '天津', '河北', '京津冀', '雄安', '张家口', '承德'],
    '长三角': ['上海', '江苏', '浙江', '安徽', '长三角', '沪苏浙皖', '南京', '杭州'],
    '珠三角': ['广东', '广州', '深圳', '珠三角', '粤港澳', '大湾区', '珠海', '东莞'],
    '东北地区': ['辽宁', '吉林', '黑龙江', '东北', '沈阳', '长春', '哈尔滨'],
    '中部地区': ['山西', '河南', '湖北', '湖南', '江西', '安徽', '中部'],
    '西部地区': ['重庆', '四川', '贵州', '云南', '西藏', '陕西', '甘肃', '青海', '宁夏', '新疆', '西部'],
    '成渝地区': ['重庆', '成都', '成渝', '四川东部', '重庆西部'],
    '长江经济带': ['长江', '长江经济带', '沿江省市', '长江流域'],
    '黄河流域': ['黄河', '黄河流域', '黄河生态经济带'],
    '全国性': ['全国', '无特定区域', '各地区', '所有地区']
}

CORE_KEYWORDS = [
    '污染防治', '排放标准', '生态保护', '碳达峰', '碳中和', '绿色转型',
    '环境治理', '执法检查', '惩处机制', '激励补贴', '双碳', '低碳转型',
    '生态文明', '可持续发展', '环境监测', '环境评估', '生态补偿',
    '绿色金融', '清洁生产', '循环经济', '无废城市', '环境修复',
    '环境规划', '环境标准', '排污许可', '环境影响评价'
]

START_YEAR = 2015
END_YEAR = 2025
START_DATE = datetime(2015, 1, 1)
END_DATE = datetime(2025, 12, 31)
MIN_ACCURACY_THRESHOLD = 0.8  # 目标准确率阈值


# -------------------------- 文件读取与内容解析 --------------------------
def read_file_content(file_path):
    """读取txt/docx/pdf内容，支持多编码和错误处理"""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.txt':
            return read_txt(file_path)
        elif ext == '.docx':
            return read_docx(file_path)
        elif ext == '.pdf':
            return read_pdf(file_path)
        else:
            print(f"不支持的文件类型: {ext}（{file_path}）")
            return ""
    except Exception as e:
        print(f"读取失败 {file_path}：{str(e)}")
        return ""


def read_txt(file_path):
    """尝试多种编码读取txt，提高兼容性"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'ansi', 'utf-16']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    print(f"无法解析编码：{file_path}")
    return ""


def read_docx(file_path):
    """读取docx文件文本内容，包括页眉页脚"""
    doc = Document(file_path)
    full_text = []

    # 读取正文
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 读取页眉
    for section in doc.sections:
        for header in section.headers:
            for para in header.paragraphs:
                full_text.append(para.text)

    # 读取页脚
    for section in doc.sections:
        for footer in section.footers:
            for para in footer.paragraphs:
                full_text.append(para.text)

    return ' '.join(full_text)


def read_pdf(file_path):
    """读取pdf文件文本内容（支持多页和加密文件）"""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            # 检查是否加密
            if reader.is_encrypted:
                try:
                    reader.decrypt('')  # 尝试空密码解密
                except:
                    print(f"加密PDF文件无法读取：{file_path}")
                    return ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
        return text
    except Exception as e:
        print(f"PDF读取错误 {file_path}：{str(e)}")
        return ""


# -------------------------- 核心信息提取与分类 --------------------------
def is_relevant(text):
    """判断文件是否与核心主题相关（至少匹配1个核心关键词）"""
    text_lower = text.lower()
    return any(kw.lower() in text_lower for kw in CORE_KEYWORDS)


def extract_date(text, file_name):
    """提取文件日期（支持多种格式，优先文本后文件名）"""
    date_patterns = [
        r'\b20\d{2}[/-]\d{1,2}[/-]\d{1,2}\b',  # YYYY-MM-DD / YYYY/MM/DD
        r'\b20\d{2}年\d{1,2}月\d{1,2}日\b',  # YYYY年MM月DD日
        r'\b20\d{2}年\d{1,2}月\b',  # YYYY年MM月
        r'\b20\d{2}年\b'  # YYYY年
    ]

    # 从文本提取
    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            return parse_date(match.group())

    # 从文件名提取
    for pattern in date_patterns:
        match = re.search(pattern, file_name)
        if match:
            return parse_date(match.group())

    return None


def parse_date(date_str):
    """解析日期字符串为datetime对象，支持部分日期"""
    try:
        if '-' in date_str or '/' in date_str:
            date_str = date_str.replace('/', '-')
            try:
                return datetime.strptime(date_str, '%Y-%m-%d')
            except:
                try:
                    return datetime.strptime(date_str, '%Y-%m')
                except:
                    return datetime.strptime(date_str, '%Y')
        elif '年' in date_str:
            try:
                return datetime.strptime(date_str, '%Y年%m月%d日')
            except:
                try:
                    return datetime.strptime(date_str, '%Y年%m月')
                except:
                    return datetime.strptime(date_str, '%Y年')
    except:
        return None
    return None


def extract_year(date_obj):
    """从日期对象提取年份（兼容None值）"""
    if isinstance(date_obj, datetime):
        return date_obj.year if (START_YEAR <= date_obj.year <= END_YEAR) else None
    return None


def classify_multiple(text, categories):
    """多标签分类，返回所有匹配的类别"""
    text_lower = text.lower()
    matched = []
    for category, keywords in categories.items():
        for kw in keywords:
            if kw.lower() in text_lower:
                matched.append(category)
                break  # 每个类别只需匹配一个关键词
    return matched if matched else ['未明确']


def get_primary_category(categories):
    """获取主要分类（第一个分类）"""
    return categories[0] if categories else '未明确'


# -------------------------- 准确性验证机制 --------------------------
def generate_detailed_summary(df, output_dir):
    """生成详细的分类统计报告，包含交叉分析和异常检测"""
    summary = []
    summary.append("===== 分类统计详细报告 =====")
    summary.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    summary.append(f"分析文件总数：{len(df)}\n")

    # 1. 环境要素 × 政策工具交叉分析
    summary.append("----- 1. 环境要素 × 政策工具分布分析 -----")
    env_tool_cross = pd.crosstab(df['主要环境要素'], df['主要政策工具'], normalize='index')
    summary.append(env_tool_cross.to_string(float_format='%.2%'))

    # 检测异常组合（占比<5%）


异常阈值 = 0.05
异常组合 = []
for env, row in env_tool_cross.iterrows():
    for tool, ratio in row.items():
        if ratio < 异常阈值 and env != '未明确' and tool != '未明确':
            异常组合.append(f"{env} + {tool}: {ratio:.2%}")

if 异常组合:
    summary.append("\n异常组合（占比<5%）：")
    summary.append("\n".join(异常组合))
    summary.append("提示：这些组合占比过低，可能存在分类逻辑问题或文献数量不足\n")
else:
    summary.append("\n未发现异常组合\n")

# 绘制热图
plt.figure(figsize=(12, 8))
sns.heatmap(env_tool_cross, annot=True, cmap="YlGnBu", fmt='.1%')
plt.title('环境要素×政策工具分布比例')
plt.tight_layout()
plt.savefig(os.path.join(output_dir, '环境要素_政策工具热图.png'))
plt.close()

# 2. 政府层级 × 区域交叉分析
summary.append("----- 2. 政府层级 × 区域分布分析 -----")
level_region_cross = pd.crosstab(df['主要政府层级'], df['主要区域'], normalize='index')
summary.append(level_region_cross.to_string(float_format='%.2%'))

# 检测异常组合
异常组合 = []
for level, row in level_region_cross.iterrows():
    for region, ratio in row.items():
        if ratio < 异常阈值 and level != '未明确' and region not in ['全国性', '未明确']:
            异常组合.append(f"{level} + {region}: {ratio:.2%}")

if 异常组合:
    summary.append("\n异常组合（占比<5%）：")
    summary.append("\n".join(异常组合))
    summary.append("提示：这些组合占比过低，可能存在区域政策覆盖不均衡问题\n")
else:
    summary.append("\n未发现异常组合\n")

# 绘制热图
plt.figure(figsize=(12, 8))
sns.heatmap(level_region_cross, annot=True, cmap="YlOrRd", fmt='.1%')
plt.title('政府层级×区域分布比例')
plt.tight_layout()
plt.savefig(os.path.join(output_dir, '政府层级_区域热图.png'))
plt.close()

# 3. 文件类型 × 政策工具交叉分析
summary.append("----- 3. 文件类型 × 政策工具分布分析 -----")
type_tool_cross = pd.crosstab(df['主要文件类型'], df['主要政策工具'], normalize='index')
summary.append(type_tool_cross.to_string(float_format='%.2%'))

# 绘制热图
plt.figure(figsize=(12, 8))
sns.heatmap(type_tool_cross, annot=True, cmap="PuBuGn", fmt='.1%')
plt.title('文件类型×政策工具分布比例')
plt.tight_layout()
plt.savefig(os.path.join(output_dir, '文件类型_政策工具热图.png'))
plt.close()

# 保存详细报告
with open(os.path.join(output_dir, '分类统计详细报告.txt'), 'w', encoding='utf-8') as f:
    f.write('\n'.join(summary))

return env_tool_cross, level_region_cross, type_tool_cross


def manual_sampling_verification(df, output_dir, sample_ratio=0.1):
    """随机抽取样本进行人工验证，并计算准确率"""
    # 创建样本验证目录
    sample_dir = os.path.join(output_dir, '人工验证样本')
    if not os.path.exists(sample_dir):
        os.makedirs(sample_dir)

    # 随机抽取样本
    sample_size = max(1, int(len(df) * sample_ratio))
    sample_df = df.sample(n=sample_size, random_state=42)

    # 保存样本数据用于人工验证
    verification_file = os.path.join(sample_dir, '待验证样本.xlsx')
    sample_df[
        ['文件路径', '文件名', '政府层级', '文件类型', '环境要素', '政策工具类型', '所属区域', '人工验证_政府层级',
         '人工验证_文件类型', '人工验证_环境要素', '人工验证_政策工具', '人工验证_区域']].to_excel(
        verification_file, index=False, engine='openpyxl')

    # 检查是否已有验证结果
    result = {
        '样本量': sample_size,
        '准确率': None,
        '各维度准确率': {},
        '验证完成': False
    }

    # 如果存在已验证的文件，计算准确率
    if os.path.exists(verification_file):
        try:
            verified_df = pd.read_excel(verification_file, engine='openpyxl')

            # 检查是否有完成的验证
            has_verification = not verified_df['人工验证_政府层级'].isna().all()

            if has_verification:
                # 过滤掉未验证的行
                valid_rows = verified_df.dropna(subset=['人工验证_政府层级'])

                # 计算各维度准确率
                dimensions = {
                    '政府层级': ('主要政府层级', '人工验证_政府层级'),
                    '文件类型': ('主要文件类型', '人工验证_文件类型'),
                    '环境要素': ('主要环境要素', '人工验证_环境要素'),
                    '政策工具': ('主要政策工具', '人工验证_政策工具'),
                    '区域': ('主要区域', '人工验证_区域')
                }

                all_preds = []
                all_trues = []

                for dim_name, (model_col, human_col) in dimensions.items():
                    # 提取主要分类进行比较
                    model_preds = valid_rows[model_col].apply(lambda x: x.split(';')[0] if ';' in x else x)
                    human_trues = valid_rows[human_col].apply(lambda x: x.split(';')[0] if ';' in x else x)

                    # 收集所有预测和真实值用于总体准确率计算
                    all_preds.extend(model_preds.tolist())
                    all_trues.extend(human_trues.tolist())

                    # 计算该维度准确率
                    acc = accuracy_score(human_trues, model_preds)
                    result['各维度准确率'][dim_name] = acc

                # 计算总体准确率
                overall_acc = accuracy_score(all_trues, all_preds)
                result['准确率'] = overall_acc
                result['验证完成'] = True

                # 生成分类报告
                class_report = classification_report(all_trues, all_preds)
                with open(os.path.join(sample_dir, '验证结果分析.txt'), 'w', encoding='utf-8') as f:
                    f.write("===== 人工验证结果分析 =====\n")
                    f.write(f"样本量: {len(valid_rows)}\n")
                    f.write(f"总体准确率: {overall_acc:.2%}\n\n")
                    f.write("各维度准确率:\n")
                    for dim, acc in result['各维度准确率'].items():
                        f.write(f"- {dim}: {acc:.2%}\n")
                    f.write("\n分类详细报告:\n")
                    f.write(class_report)

                # 绘制准确率条形图
                plt.figure(figsize=(10, 6))
                dims = list(result['各维度准确率'].keys())
                accs = list(result['各维度准确率'].values())
                plt.bar(dims, accs, color='skyblue')
                plt.axhline(y=MIN_ACCURACY_THRESHOLD, color='r', linestyle='--',
                            label=f'目标阈值 ({MIN_ACCURACY_THRESHOLD:.0%})')
                plt.ylim(0, 1.0)
                plt.ylabel('准确率')
                plt.title('各维度分类准确率')
                plt.xticks(rotation=45)
                plt.legend()
                plt.tight_layout()
                plt.savefig(os.path.join(sample_dir, '各维度准确率.png'))
                plt.close()

        except Exception as e:
            print(f"计算验证结果时出错: {str(e)}")

    return result


# -------------------------- 文献检索不足分析 --------------------------
def analyze_literature_gaps(df, cross_analyses, output_dir):
    """分析文献检索中可能存在的不足"""
    env_tool_cross, level_region_cross, type_tool_cross = cross_analyses
    gaps_report = []
    gaps_report.append("===== 文献检索不足分析报告 =====")
    gaps_report.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    # 1. 分类维度覆盖不足分析
    gaps_report.append("----- 1. 分类维度覆盖不足分析 -----")

    # 环境要素失衡分析
    env_dist = df['主要环境要素'].value_counts(normalize=True)
    env_threshold = 0.05  # 占比低于5%视为可能失衡
    underrepresented_envs = [env for env, ratio in env_dist.items()
                             if ratio < env_threshold and env != '未明确']

    if underrepresented_envs:
        gaps_report.append(f"可能存在环境要素失衡：{', '.join(underrepresented_envs)}")
        gaps_report.append("建议补充检索以下领域政策文件：")
        for env in underrepresented_envs:
            if env == '土壤':
                gaps_report.append("- 土壤污染治理、土壤修复技术相关政策")
            elif env == '碳排放':
                gaps_report.append("- 碳交易机制、双碳目标相关政策")
            else:
                gaps_report.append(f"- {env}相关的政策文件")
    else:
        gaps_report.append("环境要素分布较为均衡，未发现明显失衡")

    # 政策工具单一性分析
    tool_dist = df['主要政策工具'].value_counts(normalize=True)
    tool_threshold = 0.05
    underrepresented_tools = [tool for tool, ratio in tool_dist.items()
                              if ratio < tool_threshold and tool != '未明确']

    if underrepresented_tools:
        gaps_report.append(f"\n可能存在政策工具单一：{', '.join(underrepresented_tools)}")
        gaps_report.append("建议补充检索：")
        for tool in underrepresented_tools:
            if tool == '市场激励型':
                gaps_report.append("- 绿色金融、生态补偿、碳交易等市场机制相关政策")
            elif tool == '自愿参与型':
                gaps_report.append("- 行业自律、绿色认证、企业环境承诺相关政策")
            else:
                gaps_report.append(f"- 更多{tool}类型的政策文件")
    else:
        gaps_report.append("\n政策工具类型分布较为均衡，未发现明显单一性问题")

    # 区域覆盖不全分析
    region_dist = df['主要区域'].value_counts(normalize=True)
    region_threshold = 0.03
    underrepresented_regions = [region for region, ratio in region_dist.items()
                                if ratio < region_threshold and region not in ['未明确', '全国性']]

    if underrepresented_regions:
        gaps_report.append(f"\n可能存在区域覆盖不全：{', '.join(underrepresented_regions)}")
        gaps_report.append("建议补充检索这些区域的地方性政策文件")
    else:
        gaps_report.append("\n区域覆盖较为全面，未发现明显不足")

    # 2. 时间分布不均分析
    gaps_report.append("\n----- 2. 时间分布不均分析 -----")
    year_dist = df['年份'].value_counts().sort_index()

    if not year_dist.empty:
        gaps_report.append("各年份文件分布：")
        gaps_report.append(year_dist.to_string())

        # 分析时间段分布差异
        early_period = range(2015, 2021)
        late_period = range(2021, 2026)

        early_count = sum(year_dist.get(y, 0) for y in early_period)
        late_count = sum(year_dist.get(y, 0) for y in late_period)

        if early_count > 3 * late_count:  # 早期文献是后期3倍以上
            gaps_report.append("\n警告：2015-2020年文献远多于2021-2025年")
            gaps_report.append("建议补充检索近5年的政策文件，特别是双碳目标提出后的相关政策")

        # 分析特定年份骤减情况
        years = sorted(year_dist.index)
        for i in range(1, len(years)):
            prev_year = years[i - 1]
            curr_year = years[i]
            if year_dist[curr_year] < year_dist[prev_year] * 0.3:  # 骤减70%以上
                gaps_report.append(f"\n警告：{curr_year}年文献数量较{prev_year}年骤减超过70%")
                gaps_report.append(f"建议补充检索{curr_year}年的政策发布平台文件")
    else:
        gaps_report.append("未获取到有效的年份分布数据")

    # 3. 文件类型缺失分析
    gaps_report.append("\n----- 3. 文件类型缺失分析 -----")
    type_dist = df['主要文件类型'].value_counts(normalize=True)
    type_threshold = 0.05
    underrepresented_types = [doc_type for doc_type, ratio in type_dist.items()
                              if ratio < type_threshold and doc_type != '未明确']

    if underrepresented_types:
        gaps_report.append(f"可能存在文件类型缺失：{', '.join(underrepresented_types)}")
        gaps_report.append("建议补充检索：")
        for doc_type in underrepresented_types:
            if doc_type in ['标准', '行动方案']:
                gaps_report.append("- 行业标准数据库（如国家标准全文公开系统）和地方政府专项行动文件")
            elif doc_type in ['地方性法规', '规章']:
                gaps_report.append("- 各省市生态环境局发布的地方性文件")
            else:
                gaps_report.append(f"- 更多{doc_type}类型的政策文件")
    else:
        gaps_report.append("文件类型分布较为均衡，未发现明显缺失")

    # 4. 主题相关性不足分析
    gaps_report.append("\n----- 4. 主题相关性不足分析 -----")
    irrelevant_count = len(df[df['是否相关'] == '否'])
    irrelevant_ratio = irrelevant_count / len(df) if len(df) > 0 else 0

    if irrelevant_ratio > 0.3:  # 30%以上文件不相关
        gaps_report.append(f"警告：不相关文件占比过高（{irrelevant_ratio:.2%}）")
        gaps_report.append("建议：")
        gaps_report.append("- 补充核心主题词库，如添加行业术语：无废城市、清洁生产等")
        gaps_report.append("- 扩大检索关键词范围，如：生态环境治理、绿色发展基金等")
    else:
        gaps_report.append(f"不相关文件占比为{irrelevant_ratio:.2%}，在可接受范围内")

    # 保存分析报告
    with open(os.path.join(output_dir, '文献检索不足分析报告.txt'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(gaps_report))

    return gaps_report


# -------------------------- 总结优化步骤 --------------------------
def generate_optimization_summary(verification_result, gaps_analysis, output_dir):
    """生成总结优化步骤报告"""
    optimization = []
    optimization.append("===== 分析流程优化步骤总结 =====")
    optimization.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}\n")

    # 基础优化建议
    optimization.append("1. 分类体系优化：")
    optimization.append("- 补充「环境要素」「政策工具」分类维度，细化政府层级与文件类型")
    optimization.append("- 增加相关性判断、优先级规则和时间过滤，提升分类准确性\n")

    # 基于验证结果的优化
    optimization.append("2. 基于准确性验证的优化：")
    if verification_result['验证完成']:
        overall_acc = verification_result['准确率']
        if overall_acc < MIN_ACCURACY_THRESHOLD:
            optimization.append(
                f"- 总体准确率({overall_acc:.2%})低于目标阈值({MIN_ACCURACY_THRESHOLD:.0%})，需重点优化分类逻辑")

            # 识别准确率最低的维度
            if verification_result['各维度准确率']:
                lowest_dim = min(verification_result['各维度准确率'],
                                 key=verification_result['各维度准确率'].get)
                optimization.append(f"- 重点优化{lowest_dim}的分类标准和关键词")
        else:
            optimization.append(f"- 总体准确率({overall_acc:.2%})达到目标阈值，分类效果良好")

        # 各维度优化建议
        for dim, acc in verification_result['各维度准确率'].items():
            if acc < MIN_ACCURACY_THRESHOLD:
                optimization.append(f"- {dim}准确率({acc:.2%})不足，建议调整该维度的关键词和分类规则")
    else:
        optimization.append("- 尚未完成人工验证，请完成样本验证以获取针对性优化建议\n")

    # 基于文献缺口的优化
    optimization.append("3. 基于文献缺口的优化：")
    optimization.append("- 生成交叉统计报告，结合人工抽样验证分类效果")
    optimization.append("- 根据分类结果中占比异常的维度，针对性补充文献检索：")

    # 从缺口分析中提取具体建议
    if gaps_analysis:
        for line in gaps_analysis:
            if line.startswith("- "):
                optimization.append(f"  {line}")

    # 保存优化报告
    with open(os.path.join(output_dir, '优化步骤总结报告.txt'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(optimization))

    return optimization


# -------------------------- 主流程函数 --------------------------
def process_files(input_dir, output_dir, sample_ratio=0.1):
    """处理指定目录下的所有文件，执行完整分析流程"""
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 收集文件信息
    file_list = []
    for root, _, files in os.walk(input_dir):
        for file in files:
            file_path = os.path.join(root, file)
            file_list.append(file_path)

    # 处理每个文件并提取信息
    data = []
    for file_path in file_list:
        try:
            file_name = os.path.basename(file_path)
            content = read_file_content(file_path)
            is_rel = '是' if is_relevant(content) else '否'

            # 提取日期和年份
            date_obj = extract_date(content, file_name)
            year = extract_year(date_obj) if date_obj else None

            # 多维度分类
            levels = classify_multiple(content, LEVELS)
            doc_types = classify_multiple(content, DOC_TYPES)
            env_factors = classify_multiple(content, ENVIRONMENTAL_FACTORS)
            policy_tools = classify_multiple(content, POLICY_TOOLS)
            regions = classify_multiple(content, REGIONS)

            # 提取主要分类
            primary_level = get_primary_category(levels)
            primary_type = get_primary_category(doc_types)
            primary_env = get_primary_category(env_factors)
            primary_tool = get_primary_category(policy_tools)
            primary_region = get_primary_category(regions)

            data.append({
                '文件路径': file_path,
                '文件名': file_name,
                '是否相关': is_rel,
                '日期': date_obj.strftime('%Y-%m-%d') if date_obj else None,
                '年份': year,
                '政府层级': ';'.join(levels),
                '文件类型': ';'.join(doc_types),
                '环境要素': ';'.join(env_factors),
                '政策工具类型': ';'.join(policy_tools),
                '所属区域': ';'.join(regions),
                '主要政府层级': primary_level,
                '主要文件类型': primary_type,
                '主要环境要素': primary_env,
                '主要政策工具': primary_tool,
                '主要区域': primary_region,
                # 人工验证字段
                '人工验证_政府层级': None,
                '人工验证_文件类型': None,
                '人工验证_环境要素': None,
                '人工验证_政策工具': None,
                '人工验证_区域': None
            })
        except Exception as e:
            print(f"处理文件 {file_path} 时出错: {str(e)}")

    # 创建DataFrame
    df = pd.DataFrame(data)

    # 保存完整数据
    df.to_excel(os.path.join(output_dir, '完整分析数据.xlsx'), index=False, engine='openpyxl')

    # 生成详细分类统计报告
    cross_analyses = generate_detailed_summary(df, output_dir)

    # 进行人工抽样验证
    verification_result = manual_sampling_verification(df, output_dir, sample_ratio)

    # 分析文献检索不足
    gaps_analysis = analyze_literature_gaps(df, cross_analyses, output_dir)

    # 生成优化步骤总结
    generate_optimization_summary(verification_result, gaps_analysis, output_dir)

    print(f"分析完成！结果已保存至: {output_dir}")
    if not verification_result['验证完成']:
        print(f"请在 {os.path.join(output_dir, '人工验证样本')} 目录下完成样本验证以获取更完整的分析结果")


# -------------------------- 执行主程序 --------------------------
if __name__ == "__main__":
    # 处理命令行参数
    if len(sys.argv) > 2:
        input_directory = sys.argv[1]
        output_directory = sys.argv[2]
    else:
        # 默认目录
        input_directory = "policy_files"  # 输入文件目录
        output_directory = "analysis_results"  # 分析结果输出目录

    # 确保输入目录存在
    if not os.path.exists(input_directory):
        print(f"输入目录不存在: {input_directory}")
        print("创建默认输入目录...")
        os.makedirs(input_directory)
        print(f"请将政策文件放入 {input_directory} 目录后重新运行程序")
        sys.exit(1)

    # 执行完整分析流程
    process_files(input_directory, output_directory, sample_ratio=0.15)  # 抽取15%样本进行验证