import os
import re
import shutil
import random
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
from datetime import datetime, timedelta
from docx import Document
import PyPDF2
from collections import defaultdict, Counter
import matplotlib
import sys  # 用于处理命令行参数

matplotlib.use('Agg')  # 非交互式环境下使用
plt.rcParams["font.family"] = ["SimHei", "WenQuanYi Micro Hei", "Heiti TC"]  # 支持中文显示

# -------------------------- 多维度分类标准配置 --------------------------
LEVELS = {
    '国家层面': ['国家', '国务院', '中央', '全国', '人大', '部委', '宏观导向', '顶层设计', '国家战略',
                 '政治局', '常委会', '最高人民法院', '最高人民检察院'],
    '省层面': ['省', '自治区', '直辖市', '省人民政府', '省级', '省内', '省级规划', '省委', '省级部门'],
    '市层面': ['市', '自治州', '市人民政府', '市级', '市内', '市级规划', '市委', '市级部门'],
    '区/县层面': ['区', '县', '区政府', '县政府', '区级', '县级', '区县规划']
}

DOC_TYPES = {
    '法律': ['法', '人大颁布', '主席令', '法典', '法律'],
    '行政法规': ['条例', '行政法规', '国务院令', '实施条例'],
    '地方性法规': ['地方性法规', '省人大', '市人大', '自治区人大'],
    '规章': ['规定', '办法', '细则', '规章', '部委令', '地方政府令'],
    '规划': ['规划', '战略', '纲要', '中长期目标', '发展纲要', '专项规划'],
    '标准': ['标准', '规范', '技术要求', '排放标准', '排放限值', '技术导则', 'GB/T', 'HJ'],
    '政策文件': ['通知', '意见', '决定', '指示', '批复', '函', '政策'],
    '执法文件': ['执法', '通报', '处罚决定', '督查结果', '检查结果', '监察报告'],
    '行动方案': ['行动方案', '实施方案', '攻坚计划', '专项行动', '工作计划']
}

ENVIRONMENTAL_FACTORS = {
    '大气': ['大气', '空气', 'PM2.5', 'PM10', '雾霾', '扬尘', '挥发性有机物', 'VOCs', '二氧化硫', '氮氧化物'],
    '水': ['水', '水质', '河流', '湖泊', '黑臭水体', '污水处理', '水源地', '地下水', '地表水', '水环境'],
    '土壤': ['土壤', '土地', '重金属', '土壤修复', '耕地污染', '农用地', '建设用地', '土壤质量'],
    '碳排放': ['碳', '碳中和', '碳达峰', '双碳', '低碳', '碳交易', '碳市场', '温室气体', '碳排放权'],
    '污染源': ['污染源', '排污', '污染物', '污染源普查', '工业污染', '农业污染', '生活污染', '面源污染'],
    '生态保护': ['生态', '生态红线', '自然保护区', '生物多样性', '生态修复', '生态系统', '湿地', '森林'],
    '固废与危废': ['固体废物', '垃圾', '危废', '危险废物', '固废', '垃圾分类', '废物处理', '无废城市'],
    '噪声与振动': ['噪声', '振动', '声环境', '噪音污染', '声学'],
    '辐射': ['辐射', '放射性', '电磁辐射', '电离辐射', '核安全']
}

POLICY_TOOLS = {
    '命令控制型': ['强制', '必须', '禁止', '处罚', '关停', '限期整改', '强制执行', '标准', '许可', '审批'],
    '市场激励型': ['补贴', '奖励', '税收优惠', '碳交易', '排污权交易', '绿色金融', '环境税', '生态补偿', '绿色信贷'],
    '自愿参与型': ['自愿', '倡议', '承诺', '认证', '企业自律', '行业公约', '环境标志', '绿色认证'],
    '多元协同型': ['公众监督', 'NGO', '社会组织', '跨区域合作', '政企联动', '公众参与', '信息公开', '社会共治'],
    '技术支持型': ['技术推广', '研发', '科技', '创新', '技术标准', '示范工程', '技术改造']
}

REGIONS = {
    '京津冀': ['北京', '天津', '河北', '京津冀', '雄安', '张家口', '承德'],
    '长三角': ['上海', '江苏', '浙江', '安徽', '长三角', '沪苏浙皖', '南京', '杭州'],
    '珠三角': ['广东', '广州', '深圳', '珠三角', '粤港澳', '大湾区', '珠海', '东莞'],
    '东北地区': ['辽宁', '吉林', '黑龙江', '东北', '沈阳', '长春', '哈尔滨'],
    '中部地区': ['山西', '河南', '湖北', '湖南', '江西', '安徽', '中部'],
    '西部地区': ['重庆', '四川', '贵州', '云南', '西藏', '陕西', '甘肃', '青海', '宁夏', '新疆', '西部'],
    '成渝地区': ['重庆', '成都', '成渝', '四川东部', '重庆西部'],
    '长江经济带': ['长江', '长江经济带', '沿江省市', '长江流域'],
    '黄河流域': ['黄河', '黄河流域', '黄河生态经济带'],
    '全国性': ['全国', '无特定区域', '各地区', '所有地区']
}

CORE_KEYWORDS = [
    '污染防治', '排放标准', '生态保护', '碳达峰', '碳中和', '绿色转型',
    '环境治理', '执法检查', '惩处机制', '激励补贴', '双碳', '低碳转型',
    '生态文明', '可持续发展', '环境监测', '环境评估', '生态补偿',
    '绿色金融', '清洁生产', '循环经济', '无废城市', '环境修复',
    '环境规划', '环境标准', '排污许可', '环境影响评价'
]

START_YEAR = 2015
END_YEAR = 2025
START_DATE = datetime(2015, 1, 1)
END_DATE = datetime(2025, 12, 31)


# -------------------------- 文件读取与内容解析 --------------------------
def read_file_content(file_path):
    """读取txt/docx/pdf内容，支持多编码和错误处理"""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.txt':
            return read_txt(file_path)
        elif ext == '.docx':
            return read_docx(file_path)
        elif ext == '.pdf':
            return read_pdf(file_path)
        else:
            print(f"不支持的文件类型: {ext}（{file_path}）")
            return ""
    except Exception as e:
        print(f"读取失败 {file_path}：{str(e)}")
        return ""


def read_txt(file_path):
    """尝试多种编码读取txt，提高兼容性"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'ansi', 'utf-16']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    print(f"无法解析编码：{file_path}")
    return ""


def read_docx(file_path):
    """读取docx文件文本内容，包括页眉页脚"""
    doc = Document(file_path)
    full_text = []

    # 读取正文
    for para in doc.paragraphs:
        full_text.append(para.text)

    # 读取页眉
    for section in doc.sections:
        for header in section.headers:
            for para in header.paragraphs:
                full_text.append(para.text)

    # 读取页脚
    for section in doc.sections:
        for footer in section.footers:
            for para in footer.paragraphs:
                full_text.append(para.text)

    return ' '.join(full_text)


def read_pdf(file_path):
    """读取pdf文件文本内容（支持多页和加密文件）"""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)

            # 检查是否加密
            if reader.is_encrypted:
                try:
                    # 尝试空密码解密
                    reader.decrypt('')
                except:
                    print(f"加密PDF文件无法读取：{file_path}")
                    return ""

            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
        return text
    except Exception as e:
        print(f"PDF读取错误 {file_path}：{str(e)}")
        return ""


# -------------------------- 核心信息提取与分类 --------------------------
def is_relevant(text):
    """判断文件是否与核心主题相关（至少匹配1个核心关键词）"""
    text_lower = text.lower()
    return any(kw.lower() in text_lower for kw in CORE_KEYWORDS)


def extract_date(text, file_name):
    """提取文件日期（支持多种格式，优先文本后文件名）"""
    date_patterns = [
        r'\b20\d{2}[/-]\d{1,2}[/-]\d{1,2}\b',  # YYYY-MM-DD / YYYY/MM/DD
        r'\b20\d{2}年\d{1,2}月\d{1,2}日\b',  # YYYY年MM月DD日
        r'\b20\d{2}年\d{1,2}月\b',  # YYYY年MM月
        r'\b20\d{2}年\b'  # YYYY年
    ]

    # 从文本提取
    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            return parse_date(match.group())

    # 从文件名提取
    for pattern in date_patterns:
        match = re.search(pattern, file_name)
        if match:
            return parse_date(match.group())

    return None


def parse_date(date_str):
    """解析日期字符串为datetime对象，支持部分日期"""
    try:
        if '-' in date_str or '/' in date_str:
            date_str = date_str.replace('/', '-')
            # 尝试完整日期
            try:
                return datetime.strptime(date_str, '%Y-%m-%d')
            except:
                # 尝试年月
                try:
                    return datetime.strptime(date_str, '%Y-%m')
                except:
                    # 尝试年份
                    return datetime.strptime(date_str, '%Y')
        elif '年' in date_str:
            # 尝试完整日期
            try:
                return datetime.strptime(date_str, '%Y年%m月%d日')
            except:
                # 尝试年月
                try:
                    return datetime.strptime(date_str, '%Y年%m月')
                except:
                    # 尝试年份
                    return datetime.strptime(date_str, '%Y年')
    except:
        return None
    return None


def extract_year(date_obj):
    """从日期对象提取年份（兼容None值）"""
    if isinstance(date_obj, datetime):
        return date_obj.year if (START_YEAR <= date_obj.year <= END_YEAR) else None
    return None


def classify_by_dimension(text, dimension_dict, priority=False):
    """按维度分类（支持优先级模式和多标签）"""
    matched = []
    text_lower = text.lower()
    for category, keywords in dimension_dict.items():
        if any(kw.lower() in text_lower for kw in keywords):
            matched.append(category)
            if priority:  # 优先级模式：匹配到第一个即返回
                return matched
    return matched if matched else ['未明确']


def get_region(text, file_name):
    """区域分类（优先文件内容，补充文件名，支持多区域）"""
    combined_text = (text + " " + file_name).lower()
    regions = []
    for region, keywords in REGIONS.items():
        if any(kw.lower() in combined_text for kw in keywords):
            regions.append(region)

    # 去重并处理全国性与其他区域的关系
    regions = list(set(regions))
    if '全国性' in regions and len(regions) > 1:
        regions.remove('全国性')

    return regions if regions else ['全国性']


# -------------------------- 单文件分析与批量处理 --------------------------
def analyze_file(file_path):
    """全维度分析单个文件，返回结构化信息"""
    file_name = os.path.basename(file_path)
    text = read_file_content(file_path)
    if not text:
        return None

    # 核心信息提取
    date = extract_date(text, file_name)
    year = extract_year(date)
    in_time_range = "是" if (date and START_DATE <= date <= END_DATE) else "否"
    relevant = is_relevant(text)

    # 多维度分类
    level = classify_by_dimension(text, LEVELS, priority=True)  # 政府层级有优先级
    doc_type = classify_by_dimension(text, DOC_TYPES)
    env_factor = classify_by_dimension(text, ENVIRONMENTAL_FACTORS)
    policy_tool = classify_by_dimension(text, POLICY_TOOLS)
    region = get_region(text, file_name)
    keywords = extract_keywords(text)

    return {
        '文件名': file_name,
        '文件路径': file_path,
        '提取日期': date.strftime('%Y-%m-%d') if date else '未提取到',
        '是否在时间范围': in_time_range,
        '年份': year if year else '未明确',
        '是否相关': '是' if relevant else '否',
        '政府层级': ';'.join(level),
        '文件类型': ';'.join(doc_type),
        '环境要素': ';'.join(env_factor),
        '政策工具类型': ';'.join(policy_tool),
        '所属区域': ';'.join(region),
        '核心关键词': ';'.join(keywords)
    }


def extract_keywords(text):
    """从文本中提取核心关键词"""
    text_lower = text.lower()
    extracted = []
    for kw in CORE_KEYWORDS:
        if kw.lower() in text_lower and kw not in extracted:
            extracted.append(kw)
    return extracted if extracted else ['未匹配关键词']


def batch_analyze(input_dir):
    """批量处理目录下所有文件，返回分析结果列表"""
    results = []
    for root, _, files in os.walk(input_dir):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in ['.txt', '.docx', '.pdf']:
                file_path = os.path.join(root, file)
                info = analyze_file(file_path)
                if info:
                    results.append(info)
    return results


# -------------------------- 交叉分析与报告生成 --------------------------
def generate_cross_analysis(df, output_dir):
    """生成多维度交叉分析报告和图表"""
    # 创建输出目录
    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    # 处理多标签数据，创建主要标签列
    df['主要政府层级'] = df['政府层级'].apply(lambda x: x.split(';')[0])
    df['主要文件类型'] = df['文件类型'].apply(lambda x: x.split(';')[0])
    df['主要环境要素'] = df['环境要素'].apply(lambda x: x.split(';')[0])
    df['主要政策工具'] = df['政策工具类型'].apply(lambda x: x.split(';')[0])
    df['主要区域'] = df['所属区域'].apply(lambda x: x.split(';')[0])

    # 保存处理后的DataFrame
    df.to_excel(os.path.join(output_dir, '完整分析数据.xlsx'), index=False, engine='openpyxl')

    # 生成分析报告
    report = []
    total_files = len(df)
    relevant_files = len(df[df['是否相关'] == '是'])

    report.append("===== 环境政策文件多维度分析报告 =====")
    report.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append(f"总文件数：{total_files} 份")
    report.append(f"相关文件数：{relevant_files} 份（占比 {relevant_files / total_files:.2%}）")
    report.append(f"时间范围：{START_YEAR}年 - {END_YEAR}年\n")

    # 1. 环境要素 × 政策工具交叉分析
    report.append("----- 1. 环境要素 × 政策工具分布 -----")
    env_tool_cross = pd.crosstab(df['主要环境要素'], df['主要政策工具'], normalize='index')
    report.append(env_tool_cross.to_string())

    # 识别异常组合（占比<5%）
    异常组合 = []
    for env, row in env_tool_cross.iterrows():
        for tool, ratio in row.items():
            if ratio < 0.05 and env != '未明确' and tool != '未明确':
                异常组合.append(f"{env} + {tool}: {ratio:.2%}")

    if 异常组合:
        report.append("\n异常组合（占比<5%）：")
        report.append("\n".join(异常组合))
        report.append("提示：这些组合占比过低，可能存在分类逻辑问题或文献数量不足\n")
    else:
        report.append("\n未发现异常组合\n")

    # 绘制环境要素×政策工具热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(env_tool_cross, annot=True, cmap="YlGnBu", fmt='.2%')
    plt.title('环境要素×政策工具分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '环境要素_政策工具热图.png'))
    plt.close()

    # 2. 政府层级 × 区域交叉分析
    report.append("----- 2. 政府层级 × 区域分布 -----")
    level_region_cross = pd.crosstab(df['主要政府层级'], df['主要区域'], normalize='index')
    report.append(level_region_cross.to_string())

    # 识别异常组合
    异常组合 = []
    for level, row in level_region_cross.iterrows():
        for region, ratio in row.items():
            if ratio < 0.05 and level != '未明确' and region != '全国性' and region != '未明确':
                异常组合.append(f"{level} + {region}: {ratio:.2%}")

    if 异常组合:
        report.append("\n异常组合（占比<5%）：")
        report.append("\n".join(异常组合))
        report.append("提示：这些组合占比过低，可能存在区域政策覆盖不均衡问题\n")
    else:
        report.append("\n未发现异常组合\n")

    # 绘制政府层级×区域热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(level_region_cross, annot=True, cmap="YlOrRd", fmt='.2%')
    plt.title('政府层级×区域分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '政府层级_区域热图.png'))
    plt.close()

    # 3. 文件类型 × 政策工具交叉分析
    report.append("----- 3. 文件类型 × 政策工具分布 -----")
    type_tool_cross = pd.crosstab(df['主要文件类型'], df['主要政策工具'], normalize='index')
    report.append(type_tool_cross.to_string())

    # 绘制文件类型×政策工具热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(type_tool_cross, annot=True, cmap="PuBuGn", fmt='.2%')
    plt.title('文件类型×政策工具分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '文件类型_政策工具热图.png'))
    plt.close()

    # 4. 时间分布分析
    report.append("\n----- 4. 时间分布分析 -----")
    year_distribution = df['年份'].value_counts().sort_index()
    report.append("各年份文件数量：")
    report.append(year_distribution.to_string())

    # 绘制时间分布折线图
    plt.figure(figsize=(12, 6))
    year_distribution.plot(kind='line', marker='o')
    plt.title('文件数量年度分布')
    plt.xlabel('年份')
    plt.ylabel('文件数量')
    plt.grid(True)
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '文件数量年度分布.png'))
    plt.close()

    # 5. 各维度分布统计
    report.append("\n----- 5. 各维度分布统计 -----")

    # 政府层级分布
    level_dist = df['主要政府层级'].value_counts(normalize=True)
    report.append("政府层级分布：")
    report.append(level_dist.to_string(float_format='%.2%'))

    # 文件类型分布
    type_dist = df['主要文件类型'].value_counts(normalize=True)
    report.append("\n文件类型分布：")
    report.append(type_dist.to_string(float_format='%.2%'))

    # 环境要素分布
    env_dist = df['主要环境要素'].value_counts(normalize=True)
    report.append("\n环境要素分布：")
    report.append(env_dist.to_string(float_format='%.2%'))

    # 政策工具分布
    tool_dist = df['主要政策工具'].value_counts(normalize=True)
    report.append("\n政策工具分布：")
    report.append(tool_dist.to_string(float_format='%.2%'))

    # 区域分布
    region_dist = df['主要区域'].value_counts(normalize=True)
    report.append("\n区域分布：")
    report.append(region_dist.to_string(float_format='%.2%'))

    # 保存报告文本
    with open(os.path.join(output_dir, '分析报告.txt'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))


# -------------------------- 主程序入口 --------------------------
if __name__ == '__main__':
    # 检查命令行参数
    if len(sys.argv) != 3:
        print("用法: python test001.py <输入目录> <输出目录>")
        sys.exit(1)  # 错误退出码

    input_dir = sys.argv[1]
    output_dir = sys.argv[2]

    # 验证输入目录是否存在
    if not os.path.isdir(input_dir):
        print(f"错误：输入目录不存在 - {input_dir}")
        sys.exit(1)

    # 执行批量分析
    print(f"开始分析 {input_dir} 中的文件...")
    analysis_results = batch_analyze(input_dir)

    if not analysis_results:
        print("未找到可分析的文件（支持txt/docx/pdf）")
        sys.exit(0)

    # 转换为DataFrame并生成报告
    df = pd.DataFrame(analysis_results)
    generate_cross_analysis(df, output_dir)
    print(f"分析完成，结果已保存到 {output_dir}")