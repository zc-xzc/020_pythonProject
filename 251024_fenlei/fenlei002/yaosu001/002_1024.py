import os
import re
import shutil
import random
import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
import seaborn as sns
import subprocess
from datetime import datetime, timedelta
from docx import Document
import PyPDF2
from collections import defaultdict, Counter
import matplotlib
import sys

# -------------------------- 配置区域（请修改为实际路径） --------------------------
INPUT_DIR = r"D:\Hefei_University_of_Technology_Work\A001_2509-2605全年重点项目环境规制\C004_政策索引\北大法宝\B01部门专项层面\新建文件夹"  # 例如：r"D:\环境政策文件"
OUTPUT_DIR = r"D:\Hefei_University_of_Technology_Work\A001_2509-2605全年重点项目环境规制\C004_政策索引\北大法宝\B01部门专项层面\新建文件夹 (2)"  # 例如：r"D:\政策文件分析结果"

# 环境配置（修复字体问题）
matplotlib.use('Agg')
plt.rcParams["font.family"] = ["SimHei", "Microsoft YaHei", "Arial Unicode MS"]  # 系统兼容字体
plt.rcParams["axes.unicode_minus"] = False  # 负号显示修复

# -------------------------- 多维度分类标准配置（保留全部原有维度） --------------------------
# 政府层级分类（带优先级）
LEVELS = {
    '国家层面': ['国家', '国务院', '中央', '全国', '人大', '部委', '宏观导向', '顶层设计', '国家战略'],
    '省层面': ['省', '自治区', '直辖市', '省人民政府', '省级', '省内', '省级规划'],
    '市层面': ['市', '自治州', '市人民政府', '市级', '市内', '市级规划'],
    '区/县层面': ['区', '县', '区政府', '县政府', '区级', '县级', '区县规划']
}

# 文件类型分类
DOC_TYPES = {
    '法律': ['法', '人大颁布', '主席令', '法典', '法律'],
    '行政法规': ['条例', '行政法规', '国务院令', '实施条例'],
    '地方性法规': ['地方性法规', '省人大', '市人大', '自治区人大'],
    '规章': ['规定', '办法', '细则', '规章', '部委令', '地方政府令'],
    '规划': ['规划', '战略', '纲要', '中长期目标', '发展纲要', '专项规划'],
    '标准': ['标准', '规范', '技术要求', '排放标准', 'GB/T', 'HJ'],
    '政策文件': ['通知', '意见', '决定', '指示', '批复', '函', '政策'],
    '执法文件': ['执法', '通报', '处罚决定', '督查结果', '检查结果'],
    '行动方案': ['行动方案', '实施方案', '攻坚计划', '专项行动']
}

# 环境要素分类
ENVIRONMENTAL_FACTORS = {
    '大气': ['大气', '空气', 'PM2.5', 'PM10', '雾霾', 'VOCs', '二氧化硫', '氮氧化物'],
    '水': ['水', '水质', '河流', '湖泊', '黑臭水体', '污水处理', '水源地', '地下水'],
    '土壤': ['土壤', '土地', '重金属', '土壤修复', '耕地污染', '农用地'],
    '碳排放': ['碳', '碳中和', '碳达峰', '双碳', '低碳', '碳交易', '温室气体'],
    '污染源': ['污染源', '排污', '污染物', '工业污染', '农业污染', '生活污染'],
    '生态保护': ['生态', '生态红线', '自然保护区', '生物多样性', '生态修复'],
    '固废与危废': ['固体废物', '垃圾', '危废', '危险废物', '固废', '垃圾分类'],
    '噪声与振动': ['噪声', '振动', '声环境', '噪音污染'],
    '辐射': ['辐射', '放射性', '电磁辐射', '电离辐射']
}

# 政策工具分类
POLICY_TOOLS = {
    '命令控制型': ['强制', '必须', '禁止', '处罚', '关停', '限期整改', '许可', '审批'],
    '市场激励型': ['补贴', '奖励', '税收优惠', '碳交易', '排污权交易', '绿色金融', '生态补偿'],
    '自愿参与型': ['自愿', '倡议', '承诺', '认证', '企业自律', '行业公约'],
    '多元协同型': ['公众监督', '社会组织', '跨区域合作', '政企联动', '公众参与'],
    '技术支持型': ['技术推广', '研发', '科技', '创新', '技术标准', '示范工程']
}

# 区域分类
REGIONS = {
    '京津冀': ['北京', '天津', '河北', '京津冀', '雄安'],
    '长三角': ['上海', '江苏', '浙江', '安徽', '长三角'],
    '珠三角': ['广东', '广州', '深圳', '珠三角', '粤港澳'],
    '东北地区': ['辽宁', '吉林', '黑龙江', '东北'],
    '中部地区': ['山西', '河南', '湖北', '湖南', '江西', '安徽', '中部'],
    '西部地区': ['重庆', '四川', '贵州', '云南', '陕西', '甘肃', '西部'],
    '成渝地区': ['重庆', '成都', '成渝'],
    '长江经济带': ['长江', '长江经济带', '长江流域'],
    '黄河流域': ['黄河', '黄河流域'],
    '全国性': ['全国', '无特定区域', '各地区']
}

# 核心关键词（用于相关性判断，强化过滤逻辑）
CORE_KEYWORDS = [
    '污染防治', '排放标准', '生态保护', '碳达峰', '碳中和', '双碳',
    '绿色转型', '环境治理', '执法检查', '激励补贴', '低碳转型',
    '生态文明', '环境监测', '生态补偿', '排污许可', '环境影响评价',
    '环境保护', '环境政策', '污染控制', '生态修复', '绿色发展'
]

# 时间范围配置
START_YEAR = 2015
END_YEAR = 2025
START_DATE = datetime(2015, 1, 1)
END_DATE = datetime(2025, 12, 31)


# -------------------------- 文件读取与内容解析（保留全部格式支持） --------------------------
def read_file_content(file_path):
    """读取多种格式文件内容（txt/docx/pdf/doc），支持编码容错"""
    ext = os.path.splitext(file_path)[1].lower()
    try:
        if ext == '.txt':
            return read_txt(file_path)
        elif ext == '.docx':
            return read_docx(file_path)
        elif ext == '.pdf':
            return read_pdf(file_path)
        elif ext == '.doc':
            return read_doc(file_path)
        else:
            print(f"不支持的文件类型: {ext}（{file_path}）")
            return ""
    except Exception as e:
        print(f"读取失败 {file_path}：{str(e)}")
        return ""


def read_txt(file_path):
    """尝试多种编码读取TXT文件，提高兼容性"""
    encodings = ['utf-8', 'gbk', 'gb2312', 'ansi']
    for encoding in encodings:
        try:
            with open(file_path, 'r', encoding=encoding) as f:
                return f.read()
        except UnicodeDecodeError:
            continue
    print(f"无法解析编码：{file_path}")
    return ""


def read_docx(file_path):
    """读取DOCX文件内容（含页眉页脚）"""
    doc = Document(file_path)
    full_text = []
    # 读取正文
    for para in doc.paragraphs:
        full_text.append(para.text)
    # 读取页眉页脚
    for section in doc.sections:
        for header in section.headers:
            for para in header.paragraphs:
                full_text.append(para.text)
        for footer in section.footers:
            for para in footer.paragraphs:
                full_text.append(para.text)
    return ' '.join(full_text)


def read_pdf(file_path):
    """读取PDF文件内容（支持多页和简单加密文件）"""
    text = ""
    try:
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            # 处理加密文件
            if reader.is_encrypted:
                try:
                    reader.decrypt('')  # 尝试空密码解密
                except:
                    print(f"加密PDF文件无法读取：{file_path}")
                    return ""
            # 提取每页文本
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + " "
        return text
    except Exception as e:
        print(f"PDF读取错误 {file_path}：{str(e)}")
        return ""


def read_doc(file_path):
    """读取DOC文件内容（依赖antiword工具）"""
    try:
        result = subprocess.run(
            ['antiword', file_path],
            capture_output=True,
            text=True,
            encoding='utf-8'
        )
        return result.stdout
    except Exception as e:
        print(f"DOC文件读取失败（可忽略）：{file_path} - {str(e)}")
        return ""


# -------------------------- 核心过滤与分类逻辑（新增严格过滤） --------------------------
def is_relevant(text):
    """严格过滤无关文件：必须包含至少2个核心关键词，且总匹配次数≥3次"""
    text_lower = text.lower()
    match_count = 0
    matched_kw = set()  # 去重统计匹配的关键词

    for kw in CORE_KEYWORDS:
        kw_lower = kw.lower()
        if kw_lower in text_lower:
            matched_kw.add(kw)
            # 统计该关键词出现的次数
            match_count += text_lower.count(kw_lower)

    # 同时满足：至少2个不同关键词，且总出现次数≥3次
    return len(matched_kw) >= 2 and match_count >= 3


def extract_date(text, file_name):
    """从文本和文件名中提取日期（支持多种格式）"""
    date_patterns = [
        r'\b20\d{2}[/-]\d{1,2}[/-]\d{1,2}\b',  # YYYY-MM-DD / YYYY/MM/DD
        r'\b20\d{2}年\d{1,2}月\d{1,2}日\b',  # YYYY年MM月DD日
        r'\b20\d{2}年\d{1,2}月\b',  # YYYY年MM月
        r'\b20\d{2}年\b'  # YYYY年
    ]

    # 优先从文本提取
    for pattern in date_patterns:
        match = re.search(pattern, text)
        if match:
            return parse_date(match.group())

    # 文本无日期则从文件名提取
    for pattern in date_patterns:
        match = re.search(pattern, file_name)
        if match:
            return parse_date(match.group())

    return None


def parse_date(date_str):
    """解析日期字符串为datetime对象"""
    try:
        if '-' in date_str or '/' in date_str:
            date_str = date_str.replace('/', '-')
            try:
                return datetime.strptime(date_str, '%Y-%m-%d')
            except:
                try:
                    return datetime.strptime(date_str, '%Y-%m')
                except:
                    return datetime.strptime(date_str, '%Y')
        elif '年' in date_str:
            try:
                return datetime.strptime(date_str, '%Y年%m月%d日')
            except:
                try:
                    return datetime.strptime(date_str, '%Y年%m月')
                except:
                    return datetime.strptime(date_str, '%Y年')
    except:
        return None
    return None


def extract_year(date_obj):
    """提取年份（统一处理为字符串，避免类型错误）"""
    if isinstance(date_obj, datetime):
        year = date_obj.year
        return str(year) if (START_YEAR <= year <= END_YEAR) else '未明确'
    return '未明确'


def classify_by_dimension(text, dimension_dict, priority=False):
    """按指定维度分类（支持优先级模式）"""
    matched = []
    text_lower = text.lower()
    for category, keywords in dimension_dict.items():
        if any(kw.lower() in text_lower for kw in keywords):
            matched.append(category)
            if priority:  # 优先级模式：匹配到即终止
                return matched
    return matched if matched else ['未明确']


def get_region(text, file_name):
    """区域分类处理（解决全国性与其他区域冲突）"""
    combined_text = (text + " " + file_name).lower()
    regions = []
    for region, keywords in REGIONS.items():
        if any(kw.lower() in combined_text for kw in keywords):
            regions.append(region)
    regions = list(set(regions))
    # 全国性与其他区域共存时，移除全国性
    if '全国性' in regions and len(regions) > 1:
        regions.remove('全国性')
    return regions if regions else ['全国性']


def extract_keywords(text):
    """提取核心关键词（去重）"""
    text_lower = text.lower()
    extracted = []
    for kw in CORE_KEYWORDS:
        if kw.lower() in text_lower and kw not in extracted:
            extracted.append(kw)
    return extracted if extracted else ['未匹配关键词']


# -------------------------- 文件分类归档（保留原有归档逻辑） --------------------------
def create_category_folders(base_dir):
    """创建完整分类文件夹结构"""
    categories = {
        '按文件类型': list(DOC_TYPES.keys()) + ['未明确'],
        '按环境要素': list(ENVIRONMENTAL_FACTORS.keys()) + ['未明确'],
        '按政府层级': list(LEVELS.keys()) + ['未明确'],
        '按政策工具': list(POLICY_TOOLS.keys()) + ['未明确'],
        '按区域': list(REGIONS.keys()) + ['未明确'],
        '按年份': [str(y) for y in range(START_YEAR, END_YEAR + 1)] + ['未明确']
    }

    for main_cat, sub_cats in categories.items():
        main_path = os.path.join(base_dir, main_cat)
        os.makedirs(main_path, exist_ok=True)
        for sub_cat in sub_cats:
            sub_path = os.path.join(main_path, sub_cat)
            os.makedirs(sub_path, exist_ok=True)


def copy_to_category(file_path, analysis_result, base_output_dir):
    """根据分析结果复制文件到对应分类文件夹"""
    try:
        file_name = os.path.basename(file_path)
        year = analysis_result['年份']

        # 按文件类型复制
        for doc_type in analysis_result['文件类型'].split(';'):
            dest_path = os.path.join(base_output_dir, '按文件类型', doc_type, file_name)
            if not os.path.exists(dest_path):
                shutil.copy2(file_path, dest_path)

        # 按环境要素复制
        for env in analysis_result['环境要素'].split(';'):
            dest_path = os.path.join(base_output_dir, '按环境要素', env, file_name)
            if not os.path.exists(dest_path):
                shutil.copy2(file_path, dest_path)

        # 按政府层级复制
        for level in analysis_result['政府层级'].split(';'):
            dest_path = os.path.join(base_output_dir, '按政府层级', level, file_name)
            if not os.path.exists(dest_path):
                shutil.copy2(file_path, dest_path)

        # 按政策工具复制
        for tool in analysis_result['政策工具类型'].split(';'):
            dest_path = os.path.join(base_output_dir, '按政策工具', tool, file_name)
            if not os.path.exists(dest_path):
                shutil.copy2(file_path, dest_path)

        # 按区域复制
        for region in analysis_result['所属区域'].split(';'):
            dest_path = os.path.join(base_output_dir, '按区域', region, file_name)
            if not os.path.exists(dest_path):
                shutil.copy2(file_path, dest_path)

        # 按年份复制
        dest_path = os.path.join(base_output_dir, '按年份', year, file_name)
        if not os.path.exists(dest_path):
            shutil.copy2(file_path, dest_path)

        return True
    except Exception as e:
        print(f"文件分类失败 {file_path}：{str(e)}")
        return False


# -------------------------- 分析主函数（新增过滤步骤） --------------------------
def analyze_file(file_path):
    """全维度分析单个文件（先过滤无关文件）"""
    file_name = os.path.basename(file_path)
    text = read_file_content(file_path)

    # 空内容检查
    if not text:
        print(f"文件内容为空：{file_path}")
        return None

    # 核心过滤：先剔除无关文件
    if not is_relevant(text):
        print(f"过滤无关文件：{file_name}（核心关键词匹配不足）")
        return None

    # 核心信息提取
    date = extract_date(text, file_name)
    year = extract_year(date)  # 年份统一为字符串，避免类型错误
    in_time_range = "是" if (date and START_DATE <= date <= END_DATE) else "否"
    relevant = is_relevant(text)

    # 多维度分类
    level = classify_by_dimension(text, LEVELS, priority=True)  # 政府层级启用优先级
    doc_type = classify_by_dimension(text, DOC_TYPES)
    env_factor = classify_by_dimension(text, ENVIRONMENTAL_FACTORS)
    policy_tool = classify_by_dimension(text, POLICY_TOOLS)
    region = get_region(text, file_name)
    keywords = extract_keywords(text)

    # 结构化结果
    return {
        '文件名': file_name,
        '文件路径': file_path,
        '提取日期': date.strftime('%Y-%m-%d') if date else '未提取到',
        '是否在时间范围': in_time_range,
        '年份': year,  # 字符串类型，避免排序错误
        '是否相关': '是' if relevant else '否',
        '政府层级': ';'.join(level),
        '文件类型': ';'.join(doc_type),
        '环境要素': ';'.join(env_factor),
        '政策工具类型': ';'.join(policy_tool),
        '所属区域': ';'.join(region),
        '核心关键词': ';'.join(keywords)
    }


def batch_analyze(input_dir, output_dir):
    """批量分析目录文件（先过滤再分类）"""
    create_category_folders(output_dir)
    results = []
    supported_ext = ['.txt', '.docx', '.pdf', '.doc']  # 保留所有支持的格式

    for root, _, files in os.walk(input_dir):
        for file in files:
            ext = os.path.splitext(file)[1].lower()
            if ext in supported_ext:
                file_path = os.path.join(root, file)
                print(f"正在分析：{file_path}")
                info = analyze_file(file_path)
                if info:  # 仅保留通过过滤的文件
                    results.append(info)
                    copy_to_category(file_path, info, output_dir)

    print(f"批量分析完成，有效文件（通过过滤）：{len(results)} 个")
    return results


def generate_cross_analysis(df, output_dir):
    """生成多维度交叉分析报告（保留原有全部图表）"""
    os.makedirs(output_dir, exist_ok=True)

    # 处理多标签数据，提取主要标签
    df['主要政府层级'] = df['政府层级'].apply(lambda x: x.split(';')[0])
    df['主要文件类型'] = df['文件类型'].apply(lambda x: x.split(';')[0])
    df['主要环境要素'] = df['环境要素'].apply(lambda x: x.split(';')[0])
    df['主要政策工具'] = df['政策工具类型'].apply(lambda x: x.split(';')[0])
    df['主要区域'] = df['所属区域'].apply(lambda x: x.split(';')[0])

    # 保存完整分析数据
    df.to_excel(
        os.path.join(output_dir, '完整分析数据.xlsx'),
        index=False,
        engine='openpyxl'
    )

    # 生成分析报告文本
    report = []
    total_files = len(df)
    relevant_files = len(df[df['是否相关'] == '是'])

    report.append("===== 环境政策文件多维度分析报告 =====")
    report.append(f"生成时间：{datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    report.append(f"总文件数：{total_files} 份")
    report.append(f"相关文件数：{relevant_files} 份（占比 {relevant_files / total_files:.2%}）")
    report.append(f"时间范围：{START_YEAR}年 - {END_YEAR}年\n")

    # 1. 环境要素 × 政策工具交叉分析
    report.append("----- 1. 环境要素 × 政策工具分布 -----")
    env_tool_cross = pd.crosstab(df['主要环境要素'], df['主要政策工具'], normalize='index')
    report.append(env_tool_cross.to_string())

    # 识别异常组合
    abnormal_combinations = []
    for env, row in env_tool_cross.iterrows():
        for tool, ratio in row.items():
            if ratio < 0.05 and env != '未明确' and tool != '未明确':
                abnormal_combinations.append(f"{env} + {tool}: {ratio:.2%}")

    if abnormal_combinations:
        report.append("\n异常组合（占比<5%）：")
        report.append("\n".join(abnormal_combinations))
        report.append("提示：这些组合占比过低，可能存在分类逻辑问题或文献数量不足\n")
    else:
        report.append("\n未发现异常组合\n")

    # 绘制环境要素×政策工具热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(env_tool_cross, annot=True, cmap="YlGnBu", fmt='.2%')
    plt.title('环境要素×政策工具分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '环境要素_政策工具热图.png'))
    plt.close()

    # 2. 政府层级 × 区域交叉分析
    report.append("----- 2. 政府层级 × 区域分布 -----")
    level_region_cross = pd.crosstab(df['主要政府层级'], df['主要区域'], normalize='index')
    report.append(level_region_cross.to_string())

    # 识别异常组合
    abnormal_combinations = []
    for level, row in level_region_cross.iterrows():
        for region, ratio in row.items():
            if ratio < 0.05 and level != '未明确' and region not in ['全国性', '未明确']:
                abnormal_combinations.append(f"{level} + {region}: {ratio:.2%}")

    if abnormal_combinations:
        report.append("\n异常组合（占比<5%）：")
        report.append("\n".join(abnormal_combinations))
        report.append("提示：这些组合占比过低，可能存在区域政策覆盖不均衡问题\n")
    else:
        report.append("\n未发现异常组合\n")

    # 绘制政府层级×区域热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(level_region_cross, annot=True, cmap="YlOrRd", fmt='.2%')
    plt.title('政府层级×区域分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '政府层级_区域热图.png'))
    plt.close()

    # 3. 文件类型 × 政策工具交叉分析
    report.append("----- 3. 文件类型 × 政策工具分布 -----")
    type_tool_cross = pd.crosstab(df['主要文件类型'], df['主要政策工具'], normalize='index')
    report.append(type_tool_cross.to_string())

    # 绘制文件类型×政策工具热图
    plt.figure(figsize=(12, 8))
    sns.heatmap(type_tool_cross, annot=True, cmap="PuBuGn", fmt='.2%')
    plt.title('文件类型×政策工具分布比例')
    plt.tight_layout()
    plt.savefig(os.path.join(output_dir, '文件类型_政策工具热图.png'))
    plt.close()

    # 4. 时间分布分析（修复类型错误）
    report.append("\n----- 4. 时间分布分析 -----")
    # 筛选有效年份（排除'未明确'）
    valid_years = df[df['年份'] != '未明确']['年份']
    year_distribution = valid_years.value_counts().sort_index()
    # 补充未明确年份数量
    unknown_count = len(df[df['年份'] == '未明确'])

    report.append("各年份文件数量：")
    report.append(year_distribution.to_string())
    report.append(f"未明确年份：{unknown_count} 份")

    # 绘制时间分布折线图
    plt.figure(figsize=(12, 6))
    if not year_distribution.empty:
        year_distribution.plot(kind='line', marker='o')
        plt.title(f'文件数量年度分布（未明确：{unknown_count}份）')
        plt.xlabel('年份')
        plt.ylabel('文件数量')
        plt.grid(True)
        plt.tight_layout()
        plt.savefig(os.path.join(output_dir, '文件数量年度分布.png'))
    plt.close()

    # 5. 各维度分布统计
    report.append("\n----- 5. 各维度分布统计 -----")

    # 政府层级分布
    level_dist = df['主要政府层级'].value_counts(normalize=True)
    report.append("政府层级分布：")
    report.append(level_dist.to_string(float_format='%.2%'))

    # 文件类型分布
    type_dist = df['主要文件类型'].value_counts(normalize=True)
    report.append("\n文件类型分布：")
    report.append(type_dist.to_string(float_format='%.2%'))

    # 环境要素分布
    env_dist = df['主要环境要素'].value_counts(normalize=True)
    report.append("\n环境要素分布：")
    report.append(env_dist.to_string(float_format='%.2%'))

    # 政策工具分布
    tool_dist = df['主要政策工具'].value_counts(normalize=True)
    report.append("\n政策工具分布：")
    report.append(tool_dist.to_string(float_format='%.2%'))

    # 区域分布
    region_dist = df['主要区域'].value_counts(normalize=True)
    report.append("\n区域分布：")
    report.append(region_dist.to_string(float_format='%.2%'))

    # 保存报告文本
    with open(os.path.join(output_dir, '分析报告.txt'), 'w', encoding='utf-8') as f:
        f.write('\n'.join(report))

    print(f"交叉分析完成，结果已保存至：{output_dir}")


# -------------------------- 主函数入口（保留全部原有输出） --------------------------
def main():
    # 验证输入目录
    if not os.path.isdir(INPUT_DIR):
        print(f"错误：输入目录不存在 - {INPUT_DIR}")
        print("请修改代码中 INPUT_DIR 变量为正确的目录路径")
        return

    # 创建输出目录
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print(f"输入目录：{INPUT_DIR}")
    print(f"输出目录：{OUTPUT_DIR}\n")

    # 执行批量分析与分类（先过滤再处理）
    print("开始批量分析文件（含无关文件过滤）...")
    results = batch_analyze(INPUT_DIR, OUTPUT_DIR)

    if not results:
        print("未分析到有效文件，请检查输入目录是否包含支持的文件类型（txt/docx/pdf/doc）")
        return

    # 生成多维度分析表
    df = pd.DataFrame(results)
    table_path = os.path.join(OUTPUT_DIR, '政策文件多维度分析表.xlsx')
    df.to_excel(table_path, index=False, engine='openpyxl')
    print(f"\n多维度分析表已保存至：{table_path}")

    # 生成交叉分析报告
    cross_analysis_dir = os.path.join(OUTPUT_DIR, '交叉分析结果')
    generate_cross_analysis(df, cross_analysis_dir)

    # 生成人工抽样验证表（最多50条）
    verification_dir = os.path.join(OUTPUT_DIR, '人工验证')
    os.makedirs(verification_dir, exist_ok=True)
    sample_size = min(50, len(df))
    verification_df = df.sample(sample_size)
    verification_path = os.path.join(verification_dir, '抽样验证表.xlsx')
    verification_df.to_excel(verification_path, index=False, engine='openpyxl')
    print(f"人工抽样验证表已生成（{sample_size}条样本），保存至：{verification_path}")

    # 按时间和层级排序
    sort_dir = os.path.join(OUTPUT_DIR, '排序结果')
    os.makedirs(sort_dir, exist_ok=True)

    # 按时间排序
    time_sorted = df.sort_values(by='提取日期')
    time_sorted_path = os.path.join(sort_dir, '按时间排序.xlsx')
    time_sorted.to_excel(time_sorted_path, index=False, engine='openpyxl')

    # 按政府层级排序
    level_sorted = df.sort_values(by='政府层级')
    level_sorted_path = os.path.join(sort_dir, '按政府层级排序.xlsx')
    level_sorted.to_excel(level_sorted_path, index=False, engine='openpyxl')

    print(f"排序结果已保存至：{sort_dir}")
    print("\n所有分析流程完成！")


if __name__ == "__main__":
    main()